"""
Build the audited Word manuscript for L&O Bulletin submission.
Source: manuscripts/02_meeting_highlights_500w_v2.md
Output: submission_package/Freeman_MeetingHighlights_ASLO-SIL_2026_v2.docx

The v2 manuscript replaces the v1 .docx that carried the pre-audit claims
(Quebec City venue, 1,461 presentations, '20 years ago' historical
contrast, FT-ICR count of 2). Every claim in this build is traceable
to the underlying data inventory at output/tables/.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

OUT = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_ConferenceLandscape\submission_package")
OUT.mkdir(parents=True, exist_ok=True)

doc = Document()

# --- Document defaults ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# --- Title ---
title = doc.add_heading("Reading ourselves through the program: ASLO-SIL 2026 in numbers and verbs", level=1)

# --- Author block ---
auth = doc.add_paragraph()
r = auth.add_run("Erika C. Freeman")
r.bold = True
auth.add_run("\nLeibniz Institute of Freshwater Ecology and Inland Fisheries (IGB), Berlin, Germany")
auth.add_run("\nerika.freeman@igb-berlin.de")
auth.add_run("\nORCID: 0000-0001-7161-6038")

doc.add_paragraph()  # spacer

# --- Helper to add a paragraph with simple **bold** and *italic* markup ---
def add_styled_paragraph(text):
    p = doc.add_paragraph()
    # Process inline markup: **bold**, *italic*, "quoted" stays as plain
    # Build segments by splitting on markup tokens
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        # Add the text before the match
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith("**"):
            run = p.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = p.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p

# --- Body paragraphs (from 02_meeting_highlights_500w_v2.md, transcribed) ---

body = [
    "A conference is never only the thing that happens in the rooms. It is also a record of what a community thinks is worth gathering around. **1,455 scheduled presentations. 308 session items. About 1,400 primary presenters. About 740 institutions. Five days in Montreal.** What does that picture show?",

    "**The most-repeated session-title phrases in the 2026 program are about working across.** \"Multiple stressor(s)\" appears nine times in session titles. \"Bridging the gap\" appears eight. The exact frame *\"Towards a Convergence of Current Knowledge and Application\"* — and variants on it — appear six. Several sessions explicitly use *across ecosystems / across realms / across aquatic habitats*. Whatever else we say about the community in 2026, the program names a synthesis moment out loud.",

    "**The program structurally pairs three career-tribute sessions with six dedicated early-career sessions.** Mike Pace (EP001 \"Ecosystem Processes in a Changing World\"), Jim Cotner (EP004 \"Curating Limnology with Kindness\"), and Jim Elser (EP012 \"Ecological Stoichiometry\") are honored in the same meeting as AV001 \"Amplifying Voices in Aquatic Sciences,\" EP010A/B \"Sharing Experiences Among Early-Career Researchers,\" EC02 (ECR resilience), EC03 (cross-organisational ECR alliance), EP011 (first-timer \"How To\"), and EP006 (Raelyn Cole Editorial Fellowship retrospective). Three legacy sessions and six dedicated ECR sessions, in the same five-day program.",

    "**The SIL half is doing real work.** A keyword classifier on session and presentation titles (multi-label) finds lakes and limnology in about 28% of talks, microbial ecology in about 18%, biogeochemistry in about 17%, climate change as a forcing in about 10%, marine and oceanography in about 9%, rivers and streams in about 9%, biodiversity in about 7%, estuarine and coastal in about 6%, and food webs in about 5%. About 93% of talks carry at least one of these tags (counted as a union, not a sum). Microbial ecology and biogeochemistry tag pervasively across the program rather than living in their own dedicated tracks.",

    "**Equity is in the architecture.** Indigenous-knowledge, equity, community-led, and citizen-science framings appear in about 58 talks across 14 session items, about 4% of the program. EP013 \"Two-Eyed Seeing: Indigenous Knowledge and Western Science\" is programmed as a science session. WS05 \"Weaving Indigenous Knowledge and Western Science\" and WS08 \"From Values to Practice: Inclusive Science Spaces\" appear as dedicated workshops. SS079 \"Emerging Directions in Community-Based Water Monitoring\" treats participatory science as a methods session. AV001 features presenters from Bulacan State University (Philippines), University of Eldoret (Kenya), ILPLA (Argentina), and the Universidade Federal do Espírito Santo (Brazil), among others.",

    "**The geographic gap remains.** The US accounts for about 29% of country-detected presentations, Canada about 25%; together they are about 54% of the program. Of the top-20 represented countries, Brazil is the only one whose share of the program approximately matches its share of global renewable internal freshwater. Several countries with the largest freshwater stocks — Russia, Colombia, Indonesia, Peru, India — contribute no detected presentations.",

    "**Some methodological signatures.** AI and machine learning are tagged in about 2.4% of talks; SS070 \"Exploring the Confluence of Data, Models, and Forecasts\" names that bet explicitly. eDNA and other -omics tag about 3.2%. Mass spectrometry tags about 33 talks (about 2%), with 14 using ultra-high-resolution MS (FT-ICR, Orbitrap) and 7 naming FT-ICR specifically; most concentrate in SS050, with smaller clusters in SS049, SS041, and SS048. Long-term monitoring and time-series are the most-tagged single method, at about 9%. *Disclosure: I am giving one of those seven FT-ICR talks.*",

    "**One more thing.** The single densest day is Friday 15 May, with about 260 oral talks plus 230 posters — close to 490 presentations in one day. The program contains more parallel content than any one attendee can absorb. Pick three sessions, attend them properly, follow the bibliographies later.",

    "**The portrait the schedule gives, before the meeting actually happens.** Three legacies and six early-career sessions in the same program. Lakes, microbes, and biogeochemistry as the dominant tags. Nine sessions naming multiple stressors. Indigenous knowledge as a programmed science session. AI and -omics present as small but consistent signals. Most of the world's freshwater outside the room.",
]
for para in body:
    add_styled_paragraph(para)

doc.add_paragraph()  # spacer

# --- Figure caption ---
fc = doc.add_paragraph()
r = fc.add_run("Figure 1. ")
r.bold = True
fc.add_run('Disciplinary distribution of all 1,455 scheduled presentations at ASLO-SIL 2026, by session-level frame (keyword classifier on session names and descriptions, multi-label — columns overlap, so percentages do not sum to 100%). The three dominant tags — lakes and limnology (about 28%), microbial ecology (about 18%), and biogeochemistry (about 17%) — act as connective tissue across the program. Freshwater systems together outweigh saltwater systems. DOM chemistry is small but coherent, concentrated in the three-part SS050 symposium. Indigenous-knowledge and equity content is structurally programmed across 14 session items. About 93% of presentations carry at least one of the nine frame tags (counted as a union). Source data: scraped from the public ASLO-SIL 2026 schedule, audited 11 May 2026. Full classifier code, tag assignments, and reproducible analysis at github.com/erika-freeman/aslo2026-conference-landscape.')

doc.add_paragraph()  # spacer

# --- Footer / data citation ---
footer = doc.add_paragraph()
footer_run = footer.add_run("Data and reproducible analysis: ")
footer_run.italic = True
footer.add_run("github.com/erika-freeman/aslo2026-conference-landscape").italic = True
footer.add_run("  ·  Underlying inventory available on request.").italic = True

# --- Save ---
out_path = OUT / "Freeman_MeetingHighlights_ASLO-SIL_2026_v2.docx"
doc.save(str(out_path))
print("Saved: {}".format(out_path))

# Word count check
total_words = sum(len(p.split()) for p in body)
print("Body word count: {} (Bulletin range: 500-1500)".format(total_words))
print("\nv1 .docx preserved at: {}".format(OUT / "Freeman_MeetingHighlights_ASLO-SIL_2026.docx"))

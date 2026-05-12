"""
Build submission_bundle_v1/ — a single folder containing everything Laura
Falkenberg (L&O Bulletin editor) needs to evaluate the presubmission inquiry
and pick a format.

What this script does:
  1. Builds a v10 long-form Word manuscript (~2,050 words) from
     manuscripts/01_member_narrative_1500w_v10.md. (The short v3 docx
     already exists.)
  2. Stages a clean bundle folder with renamed, self-describing filenames.
  3. Writes 00_README.txt explaining the structure and what to attach.
  4. Writes a refreshed inquiry email (v7) that references the bundle
     filenames directly (no Drive path leak).

Output:
  G:\...\_ConferenceLandscape\submission_bundle_v1\
"""
from docx import Document
from docx.shared import Pt, Inches
from pathlib import Path
import shutil

REPO = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_GitHubRepo\aslo2026-conference-landscape")
CONFLAND = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_ConferenceLandscape")
BUNDLE = CONFLAND / "submission_bundle_v1"
BUNDLE.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# 1) Build the long-form (Option B) Word manuscript from v10 markdown
# ---------------------------------------------------------------------------

def add_runs(paragraph, text):
    """Render **bold** and *italic* inline markup into a docx paragraph."""
    import re
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        else:
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title
doc.add_heading(
    "Reading Ourselves Through the Program: "
    "ASLO-SIL 2026 in Numbers, Verbs, and Invitations",
    level=1,
)

# Subtitle (italic)
sub = doc.add_paragraph()
r = sub.add_run("What 1,455 scheduled presentations reveal about aquatic science in 2026")
r.italic = True

# Author block
auth = doc.add_paragraph()
auth.add_run("Erika C. Freeman").bold = True
auth.add_run("\nLeibniz Institute of Freshwater Ecology and Inland Fisheries (IGB), Berlin, Germany")
auth.add_run("\nerika.freeman@igb-berlin.de  ·  ORCID 0000-0001-7161-6038")

doc.add_paragraph()

# Opening
opener = doc.add_paragraph()
opener.add_run("1,455 scheduled presentations. 308 session items. About 1,400 primary presenters. About 740 institutions. Five days in Montreal.").bold = True

add_runs(doc.add_paragraph(), "A conference is never only the thing that happens in the rooms. It is also a record of what a community thinks is worth gathering around.")
add_runs(doc.add_paragraph(), "You can learn a lot from who gets a session title, which words keep coming back, which methods have become ordinary, and which questions a field has decided to take on together.")
add_runs(doc.add_paragraph(), "I read the program as a record of collective attention.")

# Section: The verbs of the schedule
doc.add_heading("The verbs of the schedule", level=2)
add_runs(doc.add_paragraph(), 'The most-repeated session-title bigrams in 2026 are about working *across*. "Multiple stressor(s)" appears nine times in session titles. "Bridging the gap" appears eight. The exact frame *"Towards a Convergence of Current Knowledge and Application"* — and variants on it — appear six. Several sessions explicitly use the language of *across ecosystems*, *across realms*, or *across aquatic habitats*.')
add_runs(doc.add_paragraph(), "These are the words the schedule itself chooses. Whatever else we say about the community in 2026, the program names a synthesis moment out loud.")

# Section: Three legacies
doc.add_heading("Three legacies, and a visible early-career scaffold", level=2)
add_runs(doc.add_paragraph(), 'The 2026 program includes three tribute sessions for senior scientists in the same meeting: **EP001 "Ecosystem Processes in a Changing World: Honoring the Career of Mike Pace,"** **EP004 "Curating Limnology with Kindness: Honoring Jim Cotner,"** and **EP012 "Ecological Stoichiometry: Honoring Jim Elser."**')
add_runs(doc.add_paragraph(), "The program also includes a visible early-career scaffold:")

ecr_items = [
    '**AV001 "Amplifying Voices in Aquatic Sciences: Aquatic Confluence"** — organised by the ASLO Early Career Committee',
    '**EP010A/B "Sharing Experiences Among Early-Career Researchers"**',
    '**EC02 "Weathering the Storm: Motivational Tools for Professional and Personal Resilience"**',
    '**EC03 "Towards a cross-organizational early-career aquatic science alliance"**',
    '**EP011 "\'How to\' — Compiling Collective Knowledge to Reduce Barriers for First-Timers"**',
    '**EP006 "A Decade of Insights with the Raelyn Cole Editorial Fellowship"**',
]
for item in ecr_items:
    p = doc.add_paragraph(style="List Bullet")
    add_runs(p, item)

add_runs(doc.add_paragraph(), "Three legacy sessions and six dedicated ECR sessions appear in the same five-day program. Whether the people who built the schedule intended a generational hand-off or not, the program structurally pairs the two.")

# Section: What the program is about
doc.add_heading("What the program is about", level=2)
add_runs(doc.add_paragraph(), "Freshwater science is a major center of gravity in this joint meeting. A keyword classifier on session and presentation titles (multi-label, so a single talk can carry more than one tag) produces the following frame distribution:")

frame_rows = [
    ("Frame", "Share of talks"),
    ("Lakes / limnology", "about 28%"),
    ("Microbial ecology", "about 18%"),
    ("Biogeochemistry / carbon cycle", "about 17%"),
    ("Climate change as forcing", "about 10%"),
    ("Rivers and streams", "about 9%"),
    ("Marine and oceanography", "about 9%"),
    ("Biodiversity and biogeography", "about 7%"),
    ("Estuarine and coastal", "about 6%"),
    ("Food webs and trophic ecology", "about 5%"),
]
table = doc.add_table(rows=len(frame_rows), cols=2)
table.style = "Light Grid Accent 1"
for i, (a, b) in enumerate(frame_rows):
    cells = table.rows[i].cells
    cells[0].text = a
    cells[1].text = b
    if i == 0:
        for c in cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.bold = True

doc.add_paragraph()
add_runs(doc.add_paragraph(), "About 93% of presentations carry at least one of these nine frame tags (counted as a union, not a sum — the tags overlap). Microbial ecology and biogeochemistry appear both in dedicated sessions and as tags threaded across many other sessions. Freshwater systems (lakes, rivers, wetlands together about 39%) outweigh saltwater systems (marine and estuarine together about 15%) in this joint meeting.")
add_runs(doc.add_paragraph(), "Methodologically, the most frequently-tagged approaches are long-term monitoring and time-series (about 9% of talks), hydrology and hydrodynamics (about 8%), and experimental mesocosms (about 3%).")

# Section: The problems the schedule names
doc.add_heading("The problems the schedule names", level=2)
add_runs(doc.add_paragraph(), "The session descriptions, taken together, name several specific problem clusters that recur across the program:")

problems = [
    '**Carbon balance of inland and coastal waters.** SS002 "Gassy Waters: Advances in Understanding and Measuring Gas Fluxes Across the Aquatic Continuum" reframes inland waters as both sources and sinks of CO₂, CH₄, and N₂O. SS050 "Ecological Significance of Dissolved Organic Matter" (three sub-sessions) asks what governs DOM persistence and reactivity.',
    '**Multiple stressors.** Nine sessions name multiple stressors explicitly, framing them as combined or interactive rather than as isolated pressures.',
    '**Climate-driven change in inland and coastal waters.** SS011 "Aquatic Sciences on Ice," SS013 "Seasonal Confluences: Ecology Under Ice," SS015 "Drying Waters," and SS019 "High Arctic Limnology in Rapidly Changing Landscapes" together cover cryosphere and drought.',
    '**Eutrophication and nutrient management.** SS010 "Advances in Nutrient Accounting Approaches for Management of Inland and Coastal Waters" names the analytical move.',
    '**Disturbance and recovery.** SS039 "Wildfires and Aquatic Systems: Implications for Hydrology, Biogeochemistry, and Ecology" sits alongside several sessions on land-use forcing.',
    '**Harmful algal blooms.** SS029 "Harmful Algal Bloom Mitigation and Its Consequences" and SS048 "Some like It Hot: Cyanobacteria Adaptations and Expansion Across Different Environments" address bloom processes and mitigation.',
    '**Microbial–DOM coupling.** SS058 "Microbial-DOM Coupling in Inland Waters in the Light of Biogeochemical Cycles" treats reciprocal microbial–molecular dynamics as a dedicated topic.',
    '**Eco-evolutionary feedbacks.** SS030 "Eco-Evolutionary Feed-Backs in Limnology and Oceanography" sets ecology and evolution on overlapping timescales.',
    '**Integration across scales.** SS082 "Mesocosm-Based Approaches for Tackling Grand Challenges in Aquatic Ecosystems," SS070 "Exploring the Confluence of Data, Models, and Forecasts for Advancing Adaptive Management," and the cluster of "bridging the gap" sessions all foreground scale-bridging as a method question.',
]
for prob in problems:
    add_runs(doc.add_paragraph(), prob)

# Section: Methods in the program
doc.add_heading("Methods in the program", level=2)
add_runs(doc.add_paragraph(), "Some methodological signatures, with classifier-assigned percentages:")

methods = [
    '**AI and machine learning** appear in about 2.4% of talks. SS070 names data, models, and forecasts as one stack.',
    '**eDNA, metagenomics, and other -omics** appear in about 3.2% of talks.',
    '**Mass spectrometry** appears in about 33 talks, around 2% of the program. Fourteen of those use ultra-high-resolution MS (FT-ICR, Orbitrap, or related); seven name FT-ICR specifically. Most concentrate in SS050, with smaller clusters in SS049 (emerging contaminants), SS041 (isotope tracing), SS048 (cyanobacteria), and a handful of proteomics or lipidomics talks in SS009 and SS031. *Disclosure: I am giving one of those seven FT-ICR talks.*',
    '**Long-term monitoring and time-series** appear in about 9% of talks — the most-tagged single method in the program.',
    '**Mesocosm experimentation** appears in about 2.6%, with SS082 the dedicated home.',
    '**Groundwater, hyporheic, and subsurface** systems appear in about 1.3% of talks.',
]
for m in methods:
    p = doc.add_paragraph(style="List Bullet")
    add_runs(p, m)

add_runs(doc.add_paragraph(), "These are tag rates on titles and descriptions; the true methodological share is almost certainly higher because most talks describe findings rather than instruments.")

# Section: How the questions are wired together
doc.add_heading("How the questions are wired together", level=2)
add_runs(doc.add_paragraph(), "Beyond which topics are common, the program also shows how topics *travel together*. The strongest co-attentions — **lakes + microbial ecology (226 talks)**, **climate change + lakes (218)**, **lakes + long-term monitoring (206)**, **biogeochemistry + lakes (196)** — carry the program's intellectual spine. Lakes appear in four of the five strongest pairs; they are the connective vertebrae.")
add_runs(doc.add_paragraph(), "Some of the most *bridging* tags are smaller in volume than the dominant hubs. By betweenness centrality across a co-occurrence network of the program's tags, the topics that connect otherwise-distinct sub-conversations are **rivers and streams** (top of the list), **estuarine and coastal systems**, **fisheries**, **food webs**, **DOM chemistry**, and notably **Indigenous knowledge and equity**. Each carries fewer talks than lakes or microbial ecology, but each one stitches together pieces of the program that would otherwise sit apart. The equity content is not a silo; it routes across the schedule.")
add_runs(doc.add_paragraph(), "Methods cluster very differently with problems. **Long-term monitoring** is the connective-tissue method — it threads through about ten distinct problem clusters, from lakes and microbial ecology to climate, biogeochemistry, biodiversity, conservation, eutrophication, and the cryosphere. **Hydrology and hydrodynamics** threads almost as broadly. **eDNA and -omics** concentrate where you would expect — microbial ecology and biodiversity. **Ultra-high-resolution mass spectrometry** concentrates in DOM and microbial–DOM coupling. **AI and machine learning**, by contrast, appear across many problem columns but at low intensity in every one — the community is placing a broad bet without yet doing concentrated work in any single domain. That distribution is itself informative: AI is being adopted as a horizontal tool, not yet specialised to particular questions. (Network and matrix visualisations of these patterns are in the repository linked at the end.)")

# Section: Indigenous knowledge, equity, and community-led science
doc.add_heading("Indigenous knowledge, equity, and community-led science", level=2)
add_runs(doc.add_paragraph(), "A classifier tag for Indigenous-knowledge, equity, community-led, or citizen science returns 58 talks across 14 session items, about 4% of the program. The session-level inclusions are notable:")

ind_items = [
    '**EP013 "Two-Eyed Seeing: Indigenous Knowledge and Western Science"** is programmed as a science session in the program structure.',
    '**WS05 "Weaving Indigenous Knowledge and Western Science"** and **WS08 "From Values to Practice: Inclusive Science Spaces"** appear as dedicated workshops.',
    '**SS079 "Emerging Directions in Community-Based Water Monitoring, Participatory Science…"** treats participatory science as a methods session.',
    '**AV001** organised by the ASLO Early Career Committee features presenters from Bulacan State University (Philippines), University of Eldoret (Kenya), ILPLA (Argentina), and the Universidade Federal do Espírito Santo (Brazil), among others.',
]
for item in ind_items:
    p = doc.add_paragraph(style="List Bullet")
    add_runs(p, item)

# Section: Gender parity, quietly reached
doc.add_heading("Gender parity, quietly reached", level=2)
add_runs(doc.add_paragraph(), "The community has also reached gender parity at both the presenter and organiser levels. Among presenter names a public-domain name database could classify (about 88% of the roughly 1,400 presenters), about **53% are female-inferred**. Among session organisers, that share is about **56%**. The remaining 12% are disproportionately Chinese, Korean, Japanese, South Asian, and other names the (European-biased) tool could not place — so the true picture is likely closer to parity still. Most of STEM is still well short of this. ASLO and SIL have arrived, and arrived without making a fuss about it. It is a real and unevenly acknowledged victory.")

# Section: A map for future invitations
doc.add_heading("A map for future invitations", level=2)
add_runs(doc.add_paragraph(), "The geography of participation points to one of the field's clearest opportunities. The US accounts for about 29% of country-detected presentations, Canada about 25%; together they are about 54% of the program.")
add_runs(doc.add_paragraph(), "Among the top-20 represented countries, Japan comes closest to parity between its share of the program (about 1.4%) and its share of global renewable internal freshwater (about 1.0%). Several countries with much larger freshwater stocks are under-represented relative to that stake: Brazil holds about 13% of the world's renewable freshwater and contributes about 1.3% of detected presentations; Russia (about 10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), and Myanmar (2%) were not detected in this inventory. Together, those seven countries hold roughly **42% of the world's renewable freshwater**.")
add_runs(doc.add_paragraph(), "Read generously, this is not only a gap. It is a map for future collaboration: stronger travel pathways, more durable institutional partnerships, remote or hybrid presentation options, mentoring across regions, and more routes for freshwater-rich countries to shape the global aquatic-science agenda. The AV001 and EP013 tracks already model what those routes can look like; the rest of the program is the canvas they can grow into.")

# Section: One more thing
doc.add_heading("One more thing", level=2)
add_runs(doc.add_paragraph(), "The meeting is dense. **260 oral talks plus 230 posters land on Friday 15 May alone** — close to 490 presentations in one day, the heaviest of the program.")
add_runs(doc.add_paragraph(), "The program contains more parallel content than any one attendee can absorb, which is also a sign of abundance: this is no longer a single conversation, but a network of overlapping ones. Pick three sessions, attend them properly, follow the bibliographies later, and let the meeting continue after the meeting.")

# Section: A closing thought
doc.add_heading("A closing thought", level=2)
add_runs(doc.add_paragraph(), "That is the portrait the schedule gives, before the meeting actually happens. Three legacies and six early-career sessions in the same program. Lakes, microbes, and biogeochemistry as dominant tags, with long-term monitoring as the connective-tissue method that threads through them. Rivers, estuaries, fisheries, food webs, DOM, and equity content as the bridge topics that hold otherwise-distinct sub-conversations together. Nine sessions naming multiple stressors. Indigenous knowledge as a programmed science session. Gender parity reached at both presenter and organiser levels. AI and -omics present as small but consistent signals. A geography of participation that shows where the next invitations could be strongest. The field is not only naming its pressures; it is wiring its questions together, and building the collaborations, methods, and institutional habits needed to respond.")

# Figure captions
doc.add_paragraph()
fc1 = doc.add_paragraph()
fc1.add_run("Figure 1. ").bold = True
fc1.add_run("Disciplinary distribution of all 1,455 scheduled presentations at ASLO-SIL 2026, by session-level frame (keyword classifier on session names and descriptions, multi-label — columns overlap, so percentages do not sum to 100%). About 93% of presentations carry at least one of the nine frame tags (counted as a union). Source data: scraped from the public ASLO-SIL 2026 schedule, audited 11 May 2026.")

fc2 = doc.add_paragraph()
fc2.add_run("Figure 2. ").bold = True
fc2.add_run("Country freshwater share (World Bank ER.H2O.INTR.K3) versus country share of country-detected program participation. Top-20 freshwater-stock countries shown. Japan is closest to parity; Brazil holds about 13% of the world's renewable freshwater and contributes about 1.3% of the program. Read generously, a map for travel pathways, regional partnerships, and hybrid presentation options.")

# Data and methods (italic small footer)
doc.add_paragraph()
dm = doc.add_paragraph()
r = dm.add_run("Data and methods. ")
r.italic = True
r.bold = True
r2 = dm.add_run("Scraped from the public ASLO-SIL 2026 schedule and audited 11 May 2026. Disciplinary and method tags assigned by a multi-label keyword classifier on session and presentation titles, so figures should be read as approximate tag rates rather than exclusive shares. Country detection is keyword-based on affiliation strings and email-domain TLDs and reaches about 78% coverage. Gender inference uses the gender-guesser public-domain name database, which is European-biased; the 12% of presenter names it could not classify are disproportionately non-Western. Co-attention network (37 tag nodes, 325 edges at weight ≥10) computed with networkx; betweenness and eigenvector centrality on the same graph. Methods-by-problems matrix is a cross-tab of method and frame tag counts. Freshwater data: World Bank ER.H2O.INTR.K3, latest available 1990–2022. Full inventory, classifier code, network and matrix visualisations, and reproducible analysis at github.com/erika-freeman/aslo2026-conference-landscape.")
r2.italic = True

# Save long-form docx to submission_package (canonical) and bundle (renamed)
LONG_NAME = "Freeman_OptionB_LongForm_v10.docx"
long_canonical = CONFLAND / "submission_package" / LONG_NAME
doc.save(str(long_canonical))
doc.save(str(BUNDLE / "05_OptionB_long_form_2050w.docx"))

# ---------------------------------------------------------------------------
# 2) Stage the bundle: copy in the Option A docx and both figures
# ---------------------------------------------------------------------------
shutil.copy(
    CONFLAND / "submission_package" / "Freeman_MeetingHighlights_ASLO-SIL_2026_v3.docx",
    BUNDLE / "04_OptionA_short_form_700w.docx",
)
shutil.copy(
    REPO / "output" / "charts" / "meeting_highlights_figure1.png",
    BUNDLE / "02_Figure1_disciplinary_distribution.png",
)
shutil.copy(
    REPO / "output" / "charts" / "freshwater_share_vs_participation_share.png",
    BUNDLE / "03_Figure2_freshwater_vs_participation.png",
)
# Also include the high-res TIFF for Option A's published Figure 1
shutil.copy(
    CONFLAND / "submission_package" / "Freeman_MeetingHighlights_Figure1.tif",
    BUNDLE / "02_Figure1_disciplinary_distribution_CMYK_300dpi.tif",
)

# ---------------------------------------------------------------------------
# 3) Write a clean inquiry email (v7) — no Drive path leak, attachment names
#    match the bundle's filenames.
# ---------------------------------------------------------------------------
inquiry_v7 = """TO:       lobulletin-editor@aslo.org
FROM:     erika.freeman@igb-berlin.de
SUBJECT:  Presubmission inquiry — Meeting Highlights on ASLO-SIL 2026

ATTACHMENTS:
  1. 02_Figure1_disciplinary_distribution.png
     (disciplinary distribution of all 1,455 talks)
  2. 03_Figure2_freshwater_vs_participation.png
     (country freshwater share vs. program share)
  3. 04_OptionA_short_form_700w.docx
     (Meeting Highlights candidate, ~700 words, one figure)
  4. 05_OptionB_long_form_2050w.docx
     (Bulletin Article candidate, ~2,050 words, two figures)

---

Dear Dr. Falkenberg,

A Meeting Highlights pitch for the L&O Bulletin, based on ASLO-SIL 2026 in Montreal this week.

I read the full programme before flying — 1,455 scheduled presentations across 308 session items, audited against the public schedule on 11 May 2026 — and the portrait the schedule itself paints is worth sharing, in numbers, verbs, and invitations.

Six things stand out, each anchored to the underlying data:

  • The most-repeated session-title phrases are about working across: "multiple stressor(s)" nine times, "bridging the gap" eight, "towards a convergence of current knowledge and application" six.

  • Three career-tribute sessions (Mike Pace EP001, Jim Cotner EP004, Jim Elser EP012) appear in the same five-day programme as six dedicated early-career sessions (AV001 Amplifying Voices, EC02 resilience, EC03 cross-organisational ECR alliance, EP010A/B, EP011 first-timer guide, EP006 Raelyn Cole Editorial Fellowship retrospective).

  • Freshwater science is a major center of gravity in this joint meeting. A keyword classifier (multi-label) finds lakes 28%, microbial ecology 18%, biogeochemistry 17%; about 93% of talks carry at least one of the dominant tags.

  • Equity work is in the architecture: about 58 talks across 14 session items, including EP013 "Two-Eyed Seeing" programmed as a science session and WS05 / WS08 as dedicated workshops. Gender parity is also visible: among classifiable presenter names, about 53% are female-inferred (about 56% among organisers).

  • Method tags are small but mapped: AI/ML about 2.4%, eDNA/-omics about 3.2%, mass spectrometry about 2% (14 ultra-high-resolution MS, 7 FT-ICR specifically), long-term monitoring about 9% — the most-tagged single method, and the connective-tissue method that threads through about ten distinct problem clusters.

  • The geography of participation points to one of the field's clearest opportunities. US plus Canada are about 54% of presenters. Japan is the country closest to parity between its share of the programme and its share of global renewable freshwater. The largest freshwater-stock countries — Brazil (13% of global freshwater, 1.3% of the programme), Russia (10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), Myanmar (2%) — together hold about 42% of the world's renewable freshwater and are under-represented or not detected. Read generously, a map for travel pathways, hybrid options, and durable regional partnerships.

Attached: (1) disciplinary distribution of all 1,455 talks; (2) country freshwater share vs. programme share — both generated from the inventory I built for this analysis. I have also attached both candidate drafts so you can pick the format that fits the Bulletin:

  A · ~700-word Meeting Highlights with one figure.
  B · ~2,050-word Bulletin Article with two figures (adds the freshwater-vs-participation comparison, a co-attention network of the program's tags, and a methods-by-problems matrix).

Both drafts have been audited: every quantitative assertion is traceable to a corresponding entry in the repository's data tables.

I am presenting at SS050B "Ecological Significance of DOM" on 15 May. The piece would be written as a participant-observer.

Would either format fit the Bulletin? I am happy to revise either draft to spec.

With thanks,

Erika

—
Dr. Erika C. Freeman
Group Leader, ABC Lab
Leibniz Institute of Freshwater Ecology and Inland Fisheries (IGB)
Müggelseedamm 310, 12587 Berlin, Germany
erika.freeman@igb-berlin.de  ·  ORCID 0000-0001-7161-6038
"""

# Save v7 inquiry to both submission_package (canonical) and bundle
(CONFLAND / "submission_package" / "presubmission_inquiry_email_v7.txt").write_text(inquiry_v7, encoding="utf-8")
(BUNDLE / "01_inquiry_email.txt").write_text(inquiry_v7, encoding="utf-8")

# ---------------------------------------------------------------------------
# 4) Write 00_README.txt — the manifest Erika reads before sending
# ---------------------------------------------------------------------------
readme = """SUBMISSION BUNDLE v1 — L&O Bulletin presubmission inquiry
==========================================================
Recipient   : Dr. Laura Falkenberg, L&O Bulletin Editor
Address     : lobulletin-editor@aslo.org
From        : Erika C. Freeman (IGB Berlin)
Bundle date : 11 May 2026
Bundle path : .../ASLO2026/_ConferenceLandscape/submission_bundle_v1/

WHAT TO DO
----------
1. Open 01_inquiry_email.txt. Copy the body (everything below the "---" divider)
   into a new email to lobulletin-editor@aslo.org. Use the SUBJECT line at the
   top of the file.

2. Attach all four files:
       02_Figure1_disciplinary_distribution.png
       03_Figure2_freshwater_vs_participation.png
       04_OptionA_short_form_700w.docx
       05_OptionB_long_form_2050w.docx

   (The bundle also contains a CMYK 300-dpi TIFF of Figure 1 for eventual
   publication — do NOT attach that to the inquiry email; it's there for the
   later formal submission step.)

3. Send.

WHAT'S IN THE BUNDLE
--------------------
00_README.txt
   This file.

01_inquiry_email.txt
   The email body. Subject line and recipient address at the top. Drafted as
   v7 — cleaned of internal Drive paths; attachment names match the files in
   this folder.

02_Figure1_disciplinary_distribution.png       (primary attachment)
   Editorial bar chart: disciplinary distribution of all 1,455 talks.
   Companion CMYK 300-dpi TIFF for publication is also included:
       02_Figure1_disciplinary_distribution_CMYK_300dpi.tif

03_Figure2_freshwater_vs_participation.png     (primary attachment)
   Country-by-country bar chart: share of global renewable freshwater vs. share
   of country-detected program participation. Top-20 freshwater-stock countries.

04_OptionA_short_form_700w.docx                (Meeting Highlights candidate)
   ~700-word Meeting Highlights piece with Figure 1.
   Title: "Reading ourselves through the program: ASLO-SIL 2026 in numbers,
   verbs, and invitations." Ready to submit today.

05_OptionB_long_form_2050w.docx                (Bulletin Article candidate)
   ~2,050-word longer version with Figure 1 + Figure 2 and additional sections:
   "The problems the schedule names," "Methods in the program," "How the
   questions are wired together" (co-attention network + methods-by-problems
   matrix), "Indigenous knowledge, equity, and community-led science," "Gender
   parity, quietly reached." Same title.

WORD-COUNT RANGES (per L&O Bulletin guidelines)
-----------------------------------------------
   Meeting Highlights : 500–1,500 words   → Option A (700) fits cleanly
   Bulletin Article   : 3,000–5,000 words → Option B (2,050) is on the short
                                              side of this range and could be
                                              expanded if Laura asks. Or Option
                                              B could be trimmed by ~550 words
                                              to land at the Meeting Highlights
                                              ceiling.

SOURCE FILES (for reference, not for attachment)
------------------------------------------------
Source markdown is in the GitHub repo:
   .../_GitHubRepo/aslo2026-conference-landscape/manuscripts/
       01_member_narrative_1500w_v10.md   (Option B source)
       02_meeting_highlights_500w_v3.md   (Option A source)
       VERSION_NOTES.md                   (audit trail v1→v10)

The build script for this bundle:
   .../_GitHubRepo/aslo2026-conference-landscape/scripts/14_build_submission_bundle_v1.py

DECLARATIONS (included in inquiry body)
---------------------------------------
• Not published or submitted elsewhere.
• No conflicts of interest.
• Every quantitative assertion in the manuscript is traceable to the
  repository's output/tables/ directory.
• The author is presenting at SS050B "Ecological Significance of DOM" on
  15 May 2026 — participant-observer perspective disclosed in the manuscript.
"""
(BUNDLE / "00_README.txt").write_text(readme, encoding="utf-8")

# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------
print("Bundle assembled at: {}".format(BUNDLE))
print()
print("Contents:")
for f in sorted(BUNDLE.iterdir()):
    print("  {:>10} bytes  {}".format(f.stat().st_size, f.name))

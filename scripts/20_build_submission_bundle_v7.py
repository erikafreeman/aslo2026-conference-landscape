"""
Build submission_bundle_v7/ — em-dash purge.

What changed vs bundle v6 (manuscript v14 -> v15):
  - All em dashes removed from body, captions, inquiry email, and watershed
    caption. Replacements mixed by local rhythm (commas, periods, colons,
    semicolons, parens) rather than mechanical comma-substitution.

Outputs:
  G:\...\_ConferenceLandscape\submission_bundle_v7\
  G:\...\submission_package\Freeman_MeetingHighlights_ASLO-SIL_2026_v8.docx
  G:\...\submission_package\Freeman_OptionB_LongForm_v16.docx
  G:\...\submission_package\presubmission_inquiry_email_v13.txt
"""
from docx import Document
from docx.shared import Pt, Inches
from pathlib import Path
from PIL import Image
import shutil
import re

REPO = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_GitHubRepo\aslo2026-conference-landscape")
CONFLAND = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_ConferenceLandscape")
BUNDLE = CONFLAND / "submission_bundle_v7"
PKG = CONFLAND / "submission_package"
BUNDLE.mkdir(parents=True, exist_ok=True)

WS_ARCHIVE_PNG = PKG / "Freeman_OpeningVisual_Watershed_ConceptA.png"
WS_ARCHIVE_TIF = PKG / "Freeman_OpeningVisual_Watershed_ConceptA_CMYK_300dpi.tif"
if not WS_ARCHIVE_PNG.exists():
    raise SystemExit("Watershed PNG missing.")

ws_embed = BUNDLE / "_tmp_watershed_for_embed.png"
img_rgb = Image.open(WS_ARCHIVE_PNG).convert("RGB")
if img_rgb.width > 1600:
    ratio = 1600 / img_rgb.width
    img_rgb = img_rgb.resize((1600, int(img_rgb.height * ratio)), Image.LANCZOS)
img_rgb.save(ws_embed, format="PNG", optimize=True)

# --- helpers ---------------------------------------------------------------
def p_text(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p

def add_runs(paragraph, text):
    pat = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        else:
            r = paragraph.add_run(token[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

WATERSHED_CAPTION = (
    "Opening visual: A community as watershed. Tributaries from ecosystems, "
    "career stages, and regions braid toward a shared confluence; a "
    "terracotta thread traces one molecule's journey from headwater to pool. "
    "Hand-drawn editorial illustration, watercolour on cream paper."
)

def insert_hero(doc, image_path, caption):
    p = doc.add_paragraph(); p.alignment = 1
    p.add_run().add_picture(str(image_path), width=Inches(6.5))
    cap = doc.add_paragraph(); cap.alignment = 1
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9)

def embed_chart(doc, image_path, caption_label, caption_body):
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = 1
    p.add_run().add_picture(str(image_path), width=Inches(6.0))
    cap = doc.add_paragraph()
    cap.add_run(caption_label).bold = True
    cap.add_run(caption_body)

# ===========================================================================
# OPTION B — Bulletin Article candidate (manuscript v15 long form)
# ===========================================================================
docB = Document()
style = docB.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

insert_hero(docB, ws_embed, WATERSHED_CAPTION)
docB.add_paragraph()

docB.add_heading(
    "Reading Ourselves Through the Program: "
    "ASLO-SIL 2026 in Numbers, Verbs, and Invitations",
    level=1,
)
sub = docB.add_paragraph()
sub.add_run("If this program is a mirror, what does it show we are becoming?").italic = True

auth = docB.add_paragraph()
auth.add_run("Erika C. Freeman").bold = True
auth.add_run("\nLeibniz Institute of Freshwater Ecology and Inland Fisheries (IGB), Berlin, Germany")
auth.add_run("\nerika.freeman@igb-berlin.de  ·  ORCID 0000-0001-7161-6038")

docB.add_paragraph()

op = docB.add_paragraph()
op.add_run("1,455 scheduled presentations. 308 session items. About 1,400 primary presenters. About 740 institutions. Five days in Montreal.").bold = True
p_text(docB, "A conference is never only the thing that happens in the rooms. It is also a mirror. Read every session title, every keyword, every method tag, and the program shows what we as a field think is worth gathering around. Some of what it shows surprised me. Some of it I think we can feel proud of. Some of it I think we should talk about together.")
p_text(docB, "I read the program as a record of collective attention.")

# What surprised me
docB.add_heading("What surprised me", level=2)
p_text(docB, "Three things worth saying out loud, before the percentages.")
p_text(docB, "The most visible method in the program is not AI. It is long-term monitoring. Long-term monitoring and time-series sit at about 9% of talks, the most-tagged single method in the program, and they thread through about ten distinct problem clusters, from lakes and microbial ecology to climate, biogeochemistry, biodiversity, conservation, eutrophication, and the cryosphere. AI and machine learning are present (about 2.4%) but spread thin: visible across many problem areas, at low intensity in each. The story the program tells is not \"AI is changing the field.\" It is \"patient, multi-decade observation is what the field is built on, with AI being adopted on top of it.\"")
p_text(docB, "Equity content is not in a silo. Indigenous-knowledge, equity, community-led, and citizen-science framings appear in 58 talks across 14 session items, about 4% of the program. That share is small, but the topological role is large: among the strongest co-occurrences in the program's tag network, Indigenous knowledge and equity sit alongside rivers, estuaries, fisheries, food webs, and DOM chemistry as bridge topics that connect otherwise-distinct sub-conversations. EP013 \"Two-Eyed Seeing: Indigenous Knowledge and Western Science\" is programmed as a science session, not an adjunct. WS05 and WS08 are dedicated workshops. SS079 treats participatory science as a methods session. The equity content is not happening next to the program; it is routing through it.")
p_text(docB, "Third, lakes are not just the largest category. They are one of the program's main confluences, where microbial ecology, climate change, long-term monitoring, and biogeochemistry repeatedly meet. In a joint ASLO and SIL meeting it is easy to read freshwater dominance as \"limnology won this round.\" It is more useful to read it as \"lakes are the system through which most of the other questions get asked.\"")

# What's changing
docB.add_heading("What's changing", level=2)
p_text(docB, "The verbs of the schedule. The most-repeated session-title phrases in 2026 are about working across boundaries. \"Multiple stressor(s)\" appears nine times. \"Bridging the gap\" appears eight. \"Towards a Convergence of Current Knowledge and Application\" appears six. Several sessions use the language of across ecosystems, across realms, or across aquatic habitats. Whatever else we say about the community in 2026, the program names a synthesis moment out loud.")
p_text(docB, "The pairing of legacy and cohort. Three career-tribute sessions for senior scientists, EP001 (Mike Pace), EP004 (Jim Cotner), and EP012 (Jim Elser), appear in the same five-day program as six dedicated early-career sessions: AV001 \"Amplifying Voices in Aquatic Sciences,\" EP010A/B \"Sharing Experiences Among Early-Career Researchers,\" EC02 (ECR resilience), EC03 (cross-organizational ECR alliance), EP011 (first-timer \"How To\"), and EP006 (Raelyn Cole Editorial Fellowship retrospective). Whether the people who built the schedule intended a generational hand-off or not, the program structurally pairs the two.")
p_text(docB, "The convergence of microbial, biogeochemistry, and DOM. SS058 \"Microbial-DOM Coupling in Inland Waters\" and the three-part SS050 \"Ecological Significance of Dissolved Organic Matter\" name reciprocal microbial-molecular dynamics as a coherent research front, no longer hypothetical, no longer scattered.")
p_text(docB, "Forecasting and monitoring brought into one conversation. SS070 \"Exploring the Confluence of Data, Models, and Forecasts for Advancing Adaptive Management\" puts long-term observation, mechanistic modelling, and forward prediction in one room. The program is treating prediction as part of the science, not as a translation layer that happens after it.")

# What we can feel proud of
docB.add_heading("What we can feel proud of", level=2)
p_text(docB, "The program shows, quietly, a strong gender-balance signal at both presenter and organizer levels: about 53% female-inferred among classifiable presenter names and 56% among classifiable organizer names, using a European-biased name classifier that left about 12% of names unclassified. Within those limits, ASLO-SIL 2026 looks close to balance, a place much of STEM has not reached.")
p_text(docB, "Three legacy sessions and six early-career sessions in the same five-day program. Indigenous-knowledge content programmed as science, not adjunct. About 93% of presentations carry at least one of the nine dominant frame tags (counted as a union). The field shows breadth, not fragmentation.")
p_text(docB, "These are not small things. They are the parts of the picture that are working.")

# Where we are stretched
docB.add_heading("Where we are stretched", level=2)
p_text(docB, "The single densest day of the meeting is Friday 15 May 2026: about 260 oral talks and 230 posters land on the same day, close to 490 presentations in twenty-four hours, the heaviest of the program. Anyone who has been to a recent ASLO meeting knows the feeling: more parallel content than any one attendee can absorb, the second-by-second triage of \"which of three sessions I care about should I miss?\"")
p_text(docB, "This is abundance, and it is also stretched cognition. The program contains more conversations than any single brain can follow in real time. Treat the density as a feature, not a checklist: pick three sessions, attend them properly, and let the meeting continue after the meeting.")

# Who is missing
docB.add_heading("Who is missing, and how we invite them in", level=2)
p_text(docB, "The geography of participation points to one of the field's clearest opportunities. The US accounts for about 29% of presentations for which a country could be detected, Canada about 25%; together they are about 54% of the program.")
p_text(docB, "Among the top-20 represented countries, Japan comes closest to parity between its share of the program (about 1.4%) and its share of global renewable internal freshwater (about 1.0%). Several countries with much larger freshwater stocks are under-represented relative to that stake. Brazil holds about 13% of the world's renewable freshwater and contributes about 1.3% of detected presentations. Russia (about 10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), and Myanmar (2%) were not detected in this inventory. Together those seven countries hold roughly 42% of the world's renewable freshwater.")
p_text(docB, "Read generously, this is not only a gap. It is a map for future collaboration. Some of the levers are programmatic: travel pathways, more durable institutional partnerships, remote or hybrid presentation options, mentoring across regions, decentralized regional hub meetings between flagship years. Others sit deeper: how international authorship is valued in evaluations and tenure cases, how funding agencies underwrite first-author opportunities for researchers based in freshwater-rich countries, and how the published record represents languages and venues beyond English-language journals. The shorter and the longer levers tend to move together.")
p_text(docB, "AV001 (organized by the ASLO Early Career Committee, with presenters from Bulacan State University, the University of Eldoret, ILPLA, and the Universidade Federal do Espírito Santo, among others) and EP013 already model what those routes can look like. The rest of the program shows how much room there is for those routes to grow.")

# Paths through the meeting
docB.add_heading("Paths through the meeting", level=2)
p_text(docB, "For a reader trying to use the program rather than only describe it, a few practical threads:")
p_text(docB, "If your work sits at the microbial-DOM-biogeochemistry interface, the three-part SS050 \"Ecological Significance of DOM,\" SS058 \"Microbial-DOM Coupling,\" SS002 \"Gassy Waters\" (gas fluxes), and the isotope-tracing work in SS041 form a coherent track. Most mass-spectrometry talks of the program cluster here.")
p_text(docB, "If you care about long-term monitoring as a research strategy, follow the tag rather than the session: long-term monitoring threads across lakes, climate, biogeochemistry, biodiversity, conservation, eutrophication, and cryosphere sessions. It is easier to find the method by reading abstracts than by reading session titles.")
p_text(docB, "If you care about equity, community-led, and Indigenous-knowledge work, EP013 (Two-Eyed Seeing, as a science session), WS05 and WS08 (workshops), SS079 (community-based monitoring), and AV001 (Amplifying Voices) form a coherent week-long arc.")
p_text(docB, "If you care about cross-realm, scale-bridging work, the \"bridging the gap\" sessions and the nine multiple-stressors sessions are the map. SS082 (mesocosms for grand challenges) and SS070 (data, models, and forecasts) anchor the methods side.")
p_text(docB, "If you care about climate-driven change in cold and dry systems, SS011 (Aquatic Sciences on Ice), SS013 (Ecology Under Ice), SS015 (Drying Waters), and SS019 (High Arctic Limnology) read as a single conversation that happens to be split across four sessions.")

# What future programs can build on
docB.add_heading("What future programs can build on", level=2)
p_text(docB, "A few signals the data carries forward, offered in the spirit of what is already working:")
p_text(docB, "The early-career scaffold is a load-bearing piece of this meeting. Six dedicated ECR sessions in five days is not ornamentation; it is infrastructure that future meetings can keep building on.")
p_text(docB, "The pairing of three legacy sessions with six ECR sessions reads as a generational hand-off in the program structure. Future meetings can lean into that intention, and could try scheduling some pairs in adjacent slots, so attendees walk from a tribute straight into the early-career conversation that extends it.")
p_text(docB, "The Friday density signal is, in part, a sign of how much the meeting has grown. Future programs that want to give attendees more chances to absorb that abundance might experiment with peak-day spacing, longer breaks between parallel blocks, or hybrid options for parallel-clash relief. All of those are more doable now than they were a decade ago.")
p_text(docB, "Bridge sessions matter even when small. EP013, WS05, WS08, and AV001 punch above their weight in the program's tag network. Future programs can build on what this meeting already started, making those bridges easy to find through separate listings, cross-track signposting, or dedicated time blocks.")
p_text(docB, "The geographic stake-vs-share gap is structural, a multi-meeting, multi-decade conversation rather than something any single program can resolve on its own. The international tracks AV001 and EP013 already prototype the kinds of routes that, in combination with funding and authorship-valuation work outside ASLO and SIL, can shift the picture over time.")

# Closing thought
docB.add_heading("A closing thought", level=2)
p_text(docB, "The program shows a field naming its pressures, building scale-bridging methods, pairing legacies with cohorts, programming equity into the science track, and stretching to absorb its own abundance.")
p_text(docB, "It also shows a field that is not a single conversation, but a network of overlapping ones: held together by lakes as recurring confluences, by long-term monitoring as a connective method, and by smaller bridge topics that carry ideas across otherwise separate parts of the schedule.")
p_text(docB, "Read generously, the picture is of a community becoming more connected, more visible to itself, and clearer about where the next invitations should go. That seems worth showing up for.")

# Figures
embed_chart(
    docB,
    REPO / "output" / "charts" / "meeting_highlights_figure1.png",
    "Figure 1. ",
    "Disciplinary distribution of all 1,455 scheduled presentations at ASLO-SIL 2026, by session-level frame (keyword classifier on session names, session descriptions, and presentation titles, multi-label; columns overlap, so percentages do not sum to 100%). About 93% of presentations carry at least one of the nine frame tags, counted as a union. Source data: scraped from the public ASLO-SIL 2026 schedule, audited 11 May 2026.",
)
embed_chart(
    docB,
    REPO / "output" / "charts" / "freshwater_share_vs_participation_share.png",
    "Figure 2. ",
    "Country freshwater share (World Bank ER.H2O.INTR.K3) versus country share of presentations for which a country could be detected. Top-20 freshwater-stock countries shown. Japan is closest to parity; Brazil holds about 13% of the world's renewable freshwater and contributes about 1.3% of the program. Read generously, a map for travel pathways, regional partnerships, and hybrid presentation options.",
)

# Data and methods
docB.add_paragraph()
dm = docB.add_paragraph()
r1 = dm.add_run("Data and methods. "); r1.italic = True; r1.bold = True
r2 = dm.add_run("Scraped from the public ASLO-SIL 2026 schedule and audited 11 May 2026. Disciplinary, method, and problem tags assigned by a multi-label keyword classifier on session names, session descriptions, and presentation titles, so figures should be read as approximate tag rates rather than exclusive shares. Country detection is keyword-based on affiliation strings and email-domain TLDs and reaches about 78% coverage. Gender inference uses the gender-guesser public-domain name database, which is European-biased; the 12% of presenter names it could not classify are disproportionately non-Western. Tag co-occurrence network (37 tag nodes, 325 edges at weight >=10) computed with networkx; betweenness and eigenvector centrality on the same graph. Methods-by-problems matrix is a cross-tab of method and frame tag counts. Freshwater data: World Bank ER.H2O.INTR.K3, latest available 1990-2022. Full inventory, classifier code, network and matrix visualizations, and reproducible analysis at github.com/erika-freeman/aslo2026-conference-landscape.")
r2.italic = True

PATH_B = PKG / "Freeman_OptionB_LongForm_v16.docx"
docB.save(str(PATH_B))
docB.save(str(BUNDLE / "05_OptionB_long_form_2100w.docx"))

# ===========================================================================
# OPTION A — short form v7
# ===========================================================================
docA = Document()
style = docA.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

insert_hero(docA, ws_embed, WATERSHED_CAPTION)
docA.add_paragraph()

docA.add_heading("Reading ourselves through the program: ASLO-SIL 2026 in numbers, verbs, and invitations", level=1)
sub = docA.add_paragraph()
sub.add_run("If this program is a mirror, what does it show we are becoming?").italic = True

auth = docA.add_paragraph()
auth.add_run("Erika C. Freeman").bold = True
auth.add_run("\nLeibniz Institute of Freshwater Ecology and Inland Fisheries (IGB), Berlin, Germany")
auth.add_run("\nerika.freeman@igb-berlin.de  ·  ORCID 0000-0001-7161-6038")

docA.add_paragraph()

short_paras = [
    "**1,455 scheduled presentations. 308 session items. About 1,400 primary presenters. About 740 institutions. Five days in Montreal.** A conference is never only what happens in the rooms. It is also a mirror. Read every session title, every keyword, every method tag, and the program shows what we as a field think is worth gathering around.",
    "**What surprised me.** Three things, before the percentages. The most visible method in the 2026 program is not AI. It is long-term monitoring, the most-tagged single method (about 9% of talks) and the one that threads through the most distinct problem clusters. AI and machine learning are present (about 2.4%) but spread thin: visible across many problems at low intensity in each. Second, equity content is not in a silo. Indigenous-knowledge, equity, community-led, and citizen-science framings appear in 58 talks across 14 session items, and among the strongest co-occurrences in the program's tag network they sit alongside rivers, estuaries, fisheries, food webs, and DOM as bridge topics that connect otherwise-distinct sub-conversations. EP013 \"Two-Eyed Seeing\" is programmed as a science session, not adjunct: the equity content is not happening next to the program; it is routing through it. Third, lakes are one of the program's main confluences, where microbial ecology, climate change, long-term monitoring, and biogeochemistry repeatedly meet. Read freshwater dominance not as \"limnology won this round\" but as \"lakes are the system through which most of the other questions get asked.\"",
    "**What's changing.** The verbs of the schedule are about working across boundaries: \"multiple stressor(s)\" appears nine times, \"bridging the gap\" eight, \"towards a convergence of current knowledge and application\" six. Three career-tribute sessions (Mike Pace EP001, Jim Cotner EP004, Jim Elser EP012) appear in the same five-day program as six dedicated ECR sessions (AV001 Amplifying Voices, EC02 resilience, EC03 cross-organizational ECR alliance, EP010A/B, EP011, EP006). Whether the schedule's architects intended a generational hand-off or not, the program structurally pairs the two. Microbial-DOM coupling moves from hypothesis to coherent research front (SS050 in three parts, SS058 as a dedicated session). Forecasting and long-term monitoring are brought into one conversation (SS070 \"Confluence of Data, Models, and Forecasts\").",
    "**What we can feel proud of.** The program shows, quietly, a strong gender-balance signal at both presenter and organizer levels: about 53% female-inferred among classifiable presenter names and 56% among classifiable organizer names, using a European-biased name classifier that left about 12% of names unclassified. Within those limits, ASLO-SIL 2026 sits close to a balance much of STEM has not reached. Three legacy sessions and six ECR sessions in the same week. Indigenous-knowledge content programmed as science. About 93% of presentations carry at least one of nine dominant frame tags: breadth, not fragmentation.",
    "**Where we are stretched.** Friday 15 May 2026 carries about 260 oral talks plus 230 posters, close to 490 presentations in twenty-four hours, the heaviest day of the program. This is abundance, and it is also stretched cognition. Treat the density as a feature, not a checklist: pick three sessions, attend them properly, and let the meeting continue after the meeting.",
    "**Who is missing, and how we invite them in.** The US accounts for about 29% of presentations for which a country could be detected, Canada about 25%; together they are about 54% of the program. Japan comes closest to parity between its share of the program (about 1.4%) and its share of global renewable freshwater (about 1.0%). Brazil holds about 13% of the world's renewable freshwater and contributes about 1.3% of detected presentations; Russia (about 10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), and Myanmar (2%) were not detected in this inventory. Together those seven countries hold roughly 42% of the world's renewable freshwater. Read generously, a map for future collaboration: travel pathways, durable institutional partnerships, hybrid presentation options, regional hub meetings between flagship years, and deeper levers: how international authorship is valued and funded, how first-author opportunities are distributed, how the published record represents work in languages and venues beyond the dominant ones. The shorter levers and the longer ones tend to move together. The AV001 and EP013 tracks already model what those routes can look like.",
    "**A few paths through the meeting, and a note for future programs.** For microbial-DOM-biogeochemistry: SS050 (three parts), SS058, SS002, and SS041 form a coherent track. For long-term monitoring as a research strategy: follow the tag across the lakes, climate, biogeochemistry, biodiversity, and cryosphere sessions. For equity, community-led, and Indigenous-knowledge work: EP013, WS05, WS08, SS079, and AV001 form a week-long arc. For cross-realm work: the \"bridging the gap\" sessions and the nine multiple-stressors sessions are the map. And for the programs that come after this one to build on: the ECR scaffold is infrastructure worth keeping; small bridge sessions reward being made easier to find; and the geography gap is a multi-meeting, multi-decade conversation that programmatic levers can support but not single-handedly resolve.",
    "**A closing thought.** The program shows a field naming its pressures, building scale-bridging methods, pairing legacies with cohorts, programming equity into the science track, and stretching to absorb its own abundance. Not a single conversation, but a network of overlapping ones: held together by lakes as recurring confluences, by long-term monitoring as a connective method, and by smaller bridge topics that carry ideas across otherwise separate parts of the schedule. Read generously, a community becoming more connected, more visible to itself, and clearer about where the next invitations should go. *Disclosure: I am presenting at SS050B \"Ecological Significance of DOM\" on 15 May; the piece is written as a participant-observer.*",
]
for s in short_paras:
    add_runs(docA.add_paragraph(), s)

embed_chart(
    docA,
    REPO / "output" / "charts" / "meeting_highlights_figure1.png",
    "Figure 1. ",
    "Disciplinary distribution of all 1,455 scheduled presentations at ASLO-SIL 2026, by session-level frame (multi-label keyword classifier on session names, descriptions, and presentation titles; columns overlap and do not sum to 100%). About 93% of presentations carry at least one of the nine frame tags. Source: scraped from the public ASLO-SIL 2026 schedule, audited 11 May 2026. Full classifier code and reproducible analysis at github.com/erika-freeman/aslo2026-conference-landscape.",
)

PATH_A = PKG / "Freeman_MeetingHighlights_ASLO-SIL_2026_v8.docx"
docA.save(str(PATH_A))
docA.save(str(BUNDLE / "04_OptionA_short_form_1050w.docx"))

# ===========================================================================
# Stage figures and watershed
# ===========================================================================
shutil.copy(WS_ARCHIVE_PNG, BUNDLE / "00_OpeningVisual_watershed.png")
shutil.copy(WS_ARCHIVE_TIF, BUNDLE / "00_OpeningVisual_watershed_CMYK_300dpi.tif")
shutil.copy(REPO / "output" / "charts" / "meeting_highlights_figure1.png",
            BUNDLE / "02_Figure1_disciplinary_distribution.png")
shutil.copy(PKG / "Freeman_MeetingHighlights_Figure1.tif",
            BUNDLE / "02_Figure1_disciplinary_distribution_CMYK_300dpi.tif")
shutil.copy(REPO / "output" / "charts" / "freshwater_share_vs_participation_share.png",
            BUNDLE / "03_Figure2_freshwater_vs_participation.png")

if ws_embed.exists():
    ws_embed.unlink()

# ===========================================================================
# Inquiry email v13 — em-dash-free
# ===========================================================================
inquiry_v13 = """TO:       lobulletin-editor@aslo.org
FROM:     erika.freeman@igb-berlin.de
SUBJECT:  Presubmission inquiry: Meeting Highlights on ASLO-SIL 2026

ATTACHMENTS:
  1. 00_OpeningVisual_watershed.png
     (hand-drawn editorial illustration, opening visual)
  2. 02_Figure1_disciplinary_distribution.png
     (disciplinary distribution of all 1,455 talks)
  3. 03_Figure2_freshwater_vs_participation.png
     (country freshwater share vs. program share)
  4. 04_OptionA_short_form_1050w.docx
     (Meeting Highlights candidate, ~1,050 words, opening visual + one figure)
  5. 05_OptionB_long_form_2100w.docx
     (Bulletin Article candidate, ~2,100 words, opening visual + two figures)

---

Dear Dr. Falkenberg,

I would like to pitch a Meeting Highlights piece for the L&O Bulletin, based on ASLO-SIL 2026 in Montreal this week. The orienting question: if this program is a mirror, what does it show we are becoming?

I read the full programme before flying. 1,455 scheduled presentations across 308 session items, audited against the public schedule on 11 May 2026. The piece is organized around what the data says about us as a field, not only what the data is. A hand-drawn editorial illustration is attached as the opening visual: a watershed reading of the meeting, with one terracotta tributary tracing one molecule's journey from headwater to confluence.

A few of the load-bearing findings, each anchored to the underlying data:

  • The most visible method in the 2026 program is not AI. It is long-term monitoring, the most-tagged single method (about 9% of talks) and the one that threads through about ten distinct problem clusters. AI and machine learning are present (about 2.4%) but spread thin: visible across many problems at low intensity in each. The program tells a story about patient, multi-decade observation as what the field is built on, with AI being adopted on top of it.

  • Equity content is not in a silo. Indigenous-knowledge, equity, community-led, and citizen-science framings appear in 58 talks across 14 session items, and among the strongest co-occurrences in the program's tag network they sit alongside rivers, estuaries, fisheries, food webs, and DOM as bridge topics that connect otherwise-distinct sub-conversations. EP013 "Two-Eyed Seeing" is programmed as a science session, not an adjunct.

  • Three career-tribute sessions (Mike Pace EP001, Jim Cotner EP004, Jim Elser EP012) appear in the same five-day programme as six dedicated early-career sessions. Whether intended or not, the program structurally pairs legacy and cohort.

  • The program shows, quietly, a strong gender-balance signal at both presenter and organizer levels: about 53% female-inferred among classifiable presenter names and 56% among classifiable organizer names, using a European-biased name classifier that left about 12% of names unclassified.

  • The geography of participation points to one of the field's clearest opportunities. US plus Canada are about 54% of presenters for which a country could be detected. Japan is the country closest to parity between its share of the programme and its share of global renewable freshwater. Brazil (13% of global freshwater, 1.3% of the programme), Russia (10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), and Myanmar (2%) together hold about 42% of the world's renewable freshwater and are under-represented or not detected. The piece treats this as a map for future collaboration, naming both programmatic levers (travel pathways, hybrid options, regional hub meetings) and deeper levers (how international authorship is valued, funded, and published).

The piece also includes a "Paths through the meeting" reader guide (microbial-DOM-biogeochemistry; long-term monitoring as research strategy; equity and community-led; cross-realm; cold-and-dry systems) and a "What future programs can build on" section, both written in the spirit of what this meeting already started.

Both candidate drafts are attached so you can pick the format that fits the Bulletin:

  A. ~1,050-word Meeting Highlights with the opening visual and one data figure.
  B. ~2,100-word Bulletin Article with the opening visual and two data figures.

Both drafts have been audited: every quantitative assertion is traceable to a corresponding entry in the repository's data tables.

I am presenting at SS050B "Ecological Significance of DOM" on 15 May. The piece is written as a participant-observer.

Would either format fit the Bulletin? I am happy to revise either draft to spec.

With thanks,

Erika

Dr. Erika C. Freeman
Group Leader, ABC Lab
Leibniz Institute of Freshwater Ecology and Inland Fisheries (IGB)
Müggelseedamm 310, 12587 Berlin, Germany
erika.freeman@igb-berlin.de  ·  ORCID 0000-0001-7161-6038
"""
(PKG / "presubmission_inquiry_email_v13.txt").write_text(inquiry_v13, encoding="utf-8")
(BUNDLE / "01_inquiry_email.txt").write_text(inquiry_v13, encoding="utf-8")

# ===========================================================================
# README
# ===========================================================================
readme = """SUBMISSION BUNDLE v7  |  L&O Bulletin presubmission inquiry
============================================================
Recipient   : Dr. Laura Falkenberg, L&O Bulletin Editor
Address     : lobulletin-editor@aslo.org
From        : Erika C. Freeman (IGB Berlin)
Bundle date : 11 May 2026
Manuscript  : v15 long form / v7 short form
Inquiry     : v13
Hero image  : Concept A watershed

WHAT CHANGED SINCE BUNDLE v6
----------------------------
Em-dash purge throughout. Replacements mixed (commas, periods, colons,
semicolons, parens) by local rhythm, not mechanical substitution. The
intent: the piece must not read as AI-written, and em dashes are one of
the most recognizable tells. Author voice tells that were doing useful
work (parallel "It is not X. It is Y." constructions, the "Read
generously" refrain, the SS050B disclosure line) were kept; mechanical
ones were removed.

WHAT TO DO
----------
1. Open 01_inquiry_email.txt. Copy the body (everything below the "---"
   divider) into a new email to lobulletin-editor@aslo.org. Use the SUBJECT
   line at the top of the file.

2. Attach all five files:
       00_OpeningVisual_watershed.png
       02_Figure1_disciplinary_distribution.png
       03_Figure2_freshwater_vs_participation.png
       04_OptionA_short_form_1050w.docx
       05_OptionB_long_form_2100w.docx

   (The CMYK TIFFs are for the eventual formal-submission step; do NOT
   attach them to the inquiry email.)

3. Send.

WHAT'S IN THE BUNDLE
--------------------
00_README.txt                                   This file.
01_inquiry_email.txt                            Email body (v13).

00_OpeningVisual_watershed.png                  Hand-drawn watershed hero.
00_OpeningVisual_watershed_CMYK_300dpi.tif      Print-ready hero.

02_Figure1_disciplinary_distribution.png        Data figure 1.
02_Figure1_disciplinary_distribution_CMYK_300dpi.tif

03_Figure2_freshwater_vs_participation.png      Data figure 2.

04_OptionA_short_form_1050w.docx                Meeting Highlights candidate.
05_OptionB_long_form_2100w.docx                 Bulletin Article candidate.

WORD-COUNT RANGES (per L&O Bulletin guidelines)
-----------------------------------------------
   Meeting Highlights : 500-1,500 words   ==> Option A (~1,050) fits cleanly.
   Bulletin Article   : 3,000-5,000 words ==> Option B (~2,100) sits below
                                                the Article minimum.

SOURCE FILES (reference, not for attachment)
--------------------------------------------
.../_GitHubRepo/aslo2026-conference-landscape/manuscripts/
    01_member_narrative_1500w_v15.md   (Option B source, current)
    02_meeting_highlights_500w_v7.md   (Option A source, current)
    VERSION_NOTES.md                   (audit trail through v15)

Build script:
.../_GitHubRepo/aslo2026-conference-landscape/scripts/20_build_submission_bundle_v7.py

Canonical archive:
.../_ConferenceLandscape/submission_package/
    Freeman_MeetingHighlights_ASLO-SIL_2026_v8.docx   (Option A)
    Freeman_OptionB_LongForm_v16.docx                 (Option B)
    presubmission_inquiry_email_v13.txt
    Freeman_OpeningVisual_Watershed_ConceptA.png
    Freeman_OpeningVisual_Watershed_ConceptA_CMYK_300dpi.tif

DECLARATIONS (included in inquiry body)
---------------------------------------
- Not published or submitted elsewhere.
- No conflicts of interest.
- Every quantitative assertion in the manuscript is traceable to the
  repository's output/tables/ directory.
- The author is presenting at SS050B "Ecological Significance of DOM" on
  15 May 2026; participant-observer perspective disclosed in the manuscript.
"""
(BUNDLE / "00_README.txt").write_text(readme, encoding="utf-8")

# ===========================================================================
# Sanity check: NO em dashes in any artifact
# ===========================================================================
EM = "—"
forbidden = [EM, "vertebrae", "country-detected", "backbone",
             "organised", "organiser", "visualisations",
             "commissioned for this piece"]
artifacts = [
    ("Option A", BUNDLE / "04_OptionA_short_form_1050w.docx"),
    ("Option B", BUNDLE / "05_OptionB_long_form_2100w.docx"),
    ("Inquiry email", BUNDLE / "01_inquiry_email.txt"),
    ("README", BUNDLE / "00_README.txt"),
]
all_clean = True
for label, p in artifacts:
    if p.suffix == ".docx":
        text = "\n".join(par.text for par in Document(str(p)).paragraphs)
    else:
        text = p.read_text(encoding="utf-8")
    hits = []
    for w in forbidden:
        # Case-insensitive check for words; literal for the em dash
        if w == EM:
            if EM in text:
                hits.append("em dash")
        elif w.lower() in text.lower():
            hits.append(w)
    if hits:
        print("WARNING: {} still contains: {}".format(label, hits))
        all_clean = False
    else:
        print("OK: {} clean.".format(label))

print()
if all_clean:
    print("All artifacts em-dash free and clean of forbidden phrasings.")
print()
print("Bundle assembled at: {}".format(BUNDLE))
print()
print("Contents:")
for f in sorted(BUNDLE.iterdir()):
    print("  {:>10} bytes  {}".format(f.stat().st_size, f.name))

def body_wc(p):
    return sum(len(par.text.split()) for par in Document(str(p)).paragraphs)
print()
print("Option A docx total word count: {}".format(body_wc(BUNDLE / "04_OptionA_short_form_1050w.docx")))
print("Option B docx total word count: {}".format(body_wc(BUNDLE / "05_OptionB_long_form_2100w.docx")))

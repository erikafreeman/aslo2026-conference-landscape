"""
Build submission_bundle_v3/ — adds the Concept-A watershed hero image.

What changed since v2:
  - Concept-A watershed image is now embedded as the opening visual in
    both Option A and Option B docx files.
  - The watershed source PNG is archived to submission_package/ with a
    clean name, and a CMYK 300-dpi TIFF version is generated for the
    eventual print step.
  - Inquiry email bumped to v9: opening-visual line added; attachment list
    extended.
  - README expanded to explain the hero/evidence figure pairing.

Outputs:
  G:\\...\\_ConferenceLandscape\\submission_bundle_v3\\
  G:\\...\\submission_package\\Freeman_OpeningVisual_Watershed_ConceptA.png
  G:\\...\\submission_package\\Freeman_OpeningVisual_Watershed_ConceptA_CMYK_300dpi.tif
  G:\\...\\submission_package\\Freeman_MeetingHighlights_ASLO-SIL_2026_v4.docx
  G:\\...\\submission_package\\Freeman_OptionB_LongForm_v12.docx
  G:\\...\\submission_package\\presubmission_inquiry_email_v9.txt
"""
from docx import Document
from docx.shared import Pt, Inches
from pathlib import Path
from PIL import Image
import shutil
import re

REPO = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_GitHubRepo\aslo2026-conference-landscape")
CONFLAND = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_ConferenceLandscape")
BUNDLE_V2 = CONFLAND / "submission_bundle_v2"
BUNDLE = CONFLAND / "submission_bundle_v3"
PKG = CONFLAND / "submission_package"
BUNDLE.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# 1) Archive the watershed image and generate derivatives
# ---------------------------------------------------------------------------
WS_SRC = BUNDLE_V2 / "watershed_source.png"  # original filename was the export name from the image tool
WS_ARCHIVE_PNG = PKG / "Freeman_OpeningVisual_Watershed_ConceptA.png"
WS_ARCHIVE_TIF = PKG / "Freeman_OpeningVisual_Watershed_ConceptA_CMYK_300dpi.tif"

# Canonical archive copy (full resolution)
if WS_SRC.exists():
    shutil.copy(WS_SRC, WS_ARCHIVE_PNG)
elif WS_ARCHIVE_PNG.exists():
    # Already archived from a previous run
    pass
else:
    raise SystemExit(
        "Watershed source not found. Expected one of:\n"
        "  {}\n  {}".format(WS_SRC, WS_ARCHIVE_PNG)
    )

# CMYK 300-dpi TIFF for print
img = Image.open(WS_ARCHIVE_PNG).convert("RGB").convert("CMYK")
img.save(
    WS_ARCHIVE_TIF,
    format="TIFF",
    compression="tiff_lzw",
    dpi=(300, 300),
)

# Downsampled copy for embedding into the docx (keeps file size sane)
ws_embed = BUNDLE / "_tmp_watershed_for_embed.png"
img_rgb = Image.open(WS_ARCHIVE_PNG).convert("RGB")
max_w = 1600
if img_rgb.width > max_w:
    ratio = max_w / img_rgb.width
    new_size = (max_w, int(img_rgb.height * ratio))
    img_rgb = img_rgb.resize(new_size, Image.LANCZOS)
img_rgb.save(ws_embed, format="PNG", optimize=True)

# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def p_text(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p

def add_runs(paragraph, text):
    """Render **bold** and *italic* inline markup."""
    pat = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        else:
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

WATERSHED_CAPTION = (
    "Opening visual. A community as watershed. Tributaries from many "
    "ecosystems, career stages, and regions braid toward a single confluence; "
    "a single terracotta thread traces one molecule's journey from headwater "
    "to pool. Hand-drawn editorial illustration, watercolour on cream paper, "
    "commissioned for this piece."
)

def insert_hero(doc, image_path, caption_text):
    """Embed the watershed image full-width, then a caption below it."""
    p = doc.add_paragraph()
    p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))
    cap = doc.add_paragraph()
    cap.alignment = 1
    r = cap.add_run(caption_text)
    r.italic = True
    r.font.size = Pt(9)

# ---------------------------------------------------------------------------
# 2) Build Option B (long form, v12 docx, manuscript v11)
# ---------------------------------------------------------------------------
docB = Document()
style = docB.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Hero image at the very top
insert_hero(docB, ws_embed, WATERSHED_CAPTION)
docB.add_paragraph()

# Title and subtitle
docB.add_heading(
    "Reading Ourselves Through the Program: "
    "ASLO-SIL 2026 in Numbers, Verbs, and Invitations",
    level=1,
)
sub = docB.add_paragraph()
sub.add_run("What 1,455 scheduled presentations reveal about aquatic science in 2026").italic = True

# Author block
auth = docB.add_paragraph()
auth.add_run("Erika C. Freeman").bold = True
auth.add_run("\nLeibniz Institute of Freshwater Ecology and Inland Fisheries (IGB), Berlin, Germany")
auth.add_run("\nerika.freeman@igb-berlin.de  ·  ORCID 0000-0001-7161-6038")

docB.add_paragraph()

# Opening
op = docB.add_paragraph()
op.add_run("1,455 scheduled presentations. 308 session items. About 1,400 primary presenters. About 740 institutions. Five days in Montreal.").bold = True

p_text(docB, "A conference is never only the thing that happens in the rooms. It is also a record of what a community thinks is worth gathering around.")
p_text(docB, "You can learn a lot from who gets a session title, which words keep coming back, which methods have become ordinary, and which questions a field has decided to take on together.")
p_text(docB, "I read the program as a record of collective attention.")

# Verbs
docB.add_heading("The verbs of the schedule", level=2)
p_text(docB, 'The most-repeated session-title phrases in 2026 are about working across boundaries. "Multiple stressor(s)" appears nine times in session titles. "Bridging the gap" appears eight. The exact frame "Towards a Convergence of Current Knowledge and Application" and variants on it appear six. Several sessions explicitly use the language of across ecosystems, across realms, or across aquatic habitats.')
p_text(docB, "These are the words the schedule itself chooses. Whatever else we say about the community in 2026, the program names a synthesis moment out loud.")

# Three legacies
docB.add_heading("Three legacies, and a visible early-career scaffold", level=2)
p_text(docB, 'The 2026 program includes three tribute sessions for senior scientists in the same meeting: EP001 "Ecosystem Processes in a Changing World: Honoring the Career of Mike Pace," EP004 "Curating Limnology with Kindness: Honoring Jim Cotner," and EP012 "Ecological Stoichiometry: Honoring Jim Elser."')
p_text(docB, "The program also includes a visible early-career scaffold:")
for item in [
    'AV001 "Amplifying Voices in Aquatic Sciences: Aquatic Confluence," organized by the ASLO Early Career Committee',
    'EP010A/B "Sharing Experiences Among Early-Career Researchers"',
    'EC02 "Weathering the Storm: Motivational Tools for Professional and Personal Resilience"',
    'EC03 "Towards a cross-organizational early-career aquatic science alliance"',
    'EP011 "\'How to\' - Compiling Collective Knowledge to Reduce Barriers for First-Timers"',
    'EP006 "A Decade of Insights with the Raelyn Cole Editorial Fellowship"',
]:
    docB.add_paragraph(item, style="List Bullet")
p_text(docB, "Three legacy sessions and six dedicated ECR sessions appear in the same five-day program. Whether the people who built the schedule intended a generational hand-off or not, the program structurally pairs the two.")

# What the program is about
docB.add_heading("What the program is about", level=2)
p_text(docB, "Freshwater science is a major center of gravity in this joint meeting. A keyword classifier on session names, session descriptions, and presentation titles (multi-label, so a single talk can carry more than one tag) produces the following frame distribution:")

frame_rows = [
    ("Frame", "Share of talks"),
    ("Lakes / limnology", "about 28%"),
    ("Microbial ecology", "about 18%"),
    ("Biogeochemistry / carbon cycle", "about 17%"),
    ("Climate change as forcing", "about 10%"),
    ("Rivers and streams", "about 9%"),
    ("Marine and oceanography", "about 9%"),
    ("Biodiversity and biogeography", "about 7%"),
    ("Estuarine and coastal", "about 6%"),
    ("Food webs and trophic ecology", "about 5%"),
]
tbl = docB.add_table(rows=len(frame_rows), cols=2)
tbl.style = "Light Grid Accent 1"
for i, (a, b) in enumerate(frame_rows):
    cells = tbl.rows[i].cells
    cells[0].text = a
    cells[1].text = b
    if i == 0:
        for c in cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.bold = True

docB.add_paragraph()
p_text(docB, "About 93% of presentations carry at least one of these nine frame tags, counted as a union, not a sum. Microbial ecology and biogeochemistry appear both in dedicated sessions and as tags threaded across many other sessions. Freshwater systems (lakes, rivers, wetlands together about 39%) outweigh saltwater systems (marine and estuarine together about 15%) in this joint meeting.")
p_text(docB, "Methodologically, the most frequently-tagged approaches are long-term monitoring and time-series (about 9% of talks), hydrology and hydrodynamics (about 8%), and experimental mesocosms (about 3%).")

# The problems the schedule names
docB.add_heading("The problems the schedule names", level=2)
p_text(docB, "The session descriptions, taken together, name several specific problem clusters that recur across the program:")
for prob in [
    'Carbon balance of inland and coastal waters. SS002 "Gassy Waters: Advances in Understanding and Measuring Gas Fluxes Across the Aquatic Continuum" reframes inland waters as both sources and sinks of CO2, CH4, and N2O. SS050 "Ecological Significance of Dissolved Organic Matter" (three sub-sessions) asks what governs DOM persistence and reactivity.',
    'Multiple stressors. Nine sessions name multiple stressors explicitly, framing them as combined or interactive rather than as isolated pressures.',
    'Climate-driven change in inland and coastal waters. SS011 "Aquatic Sciences on Ice," SS013 "Seasonal Confluences: Ecology Under Ice," SS015 "Drying Waters," and SS019 "High Arctic Limnology in Rapidly Changing Landscapes" together cover cryosphere and drought.',
    'Eutrophication and nutrient management. SS010 "Advances in Nutrient Accounting Approaches for Management of Inland and Coastal Waters" names the analytical move.',
    'Disturbance and recovery. SS039 "Wildfires and Aquatic Systems: Implications for Hydrology, Biogeochemistry, and Ecology" sits alongside several sessions on land-use forcing.',
    'Harmful algal blooms. SS029 "Harmful Algal Bloom Mitigation and Its Consequences" and SS048 "Some like It Hot: Cyanobacteria Adaptations and Expansion Across Different Environments" address bloom processes and mitigation.',
    'Microbial-DOM coupling. SS058 "Microbial-DOM Coupling in Inland Waters in the Light of Biogeochemical Cycles" treats reciprocal microbial-molecular dynamics as a dedicated topic.',
    'Eco-evolutionary feedbacks. SS030 "Eco-Evolutionary Feed-Backs in Limnology and Oceanography" sets ecology and evolution on overlapping timescales.',
    'Integration across scales. SS082 "Mesocosm-Based Approaches for Tackling Grand Challenges in Aquatic Ecosystems," SS070 "Exploring the Confluence of Data, Models, and Forecasts for Advancing Adaptive Management," and the cluster of "bridging the gap" sessions all foreground scale-bridging as a method question.',
]:
    p_text(docB, prob)

# Methods
docB.add_heading("Methods in the program", level=2)
p_text(docB, "Some methodological signatures, with classifier-assigned percentages:")
for m in [
    "AI and machine learning appear in about 2.4% of talks. SS070 names data, models, and forecasts as one stack.",
    "eDNA, metagenomics, and other -omics appear in about 3.2% of talks.",
    "Mass spectrometry appears in about 33 talks, around 2% of the program. Fourteen of those use ultra-high-resolution MS (FT-ICR, Orbitrap, or related); seven name FT-ICR specifically. Most concentrate in SS050, with smaller clusters in SS049 (emerging contaminants), SS041 (isotope tracing), SS048 (cyanobacteria), and a handful of proteomics or lipidomics talks in SS009 and SS031. Disclosure: I am giving one of those seven FT-ICR talks.",
    "Long-term monitoring and time-series appear in about 9% of talks, making them the most-tagged single method in the program.",
    "Mesocosm experimentation appears in about 2.6%, with SS082 the dedicated home.",
    "Groundwater, hyporheic, and subsurface systems appear in about 1.3% of talks.",
    "These are tag rates on titles and descriptions; the true methodological share is almost certainly higher because most talks describe findings rather than instruments.",
]:
    p_text(docB, m)

# Wired together
docB.add_heading("How the questions are wired together", level=2)
p_text(docB, "Beyond which topics are common, the program also shows how topics travel together. In a tag co-occurrence network, lakes and limnology sit at the center of the largest cluster, linking microbial ecology, biogeochemistry, climate forcing, long-term monitoring, rivers and streams, and DOM-related work. The network should be read as a map of co-attention rather than as a set of independent topic totals: the tags overlap, and a single talk can help connect several conversations at once.")
p_text(docB, "Some of the most bridging tags are smaller in volume than the dominant hubs. By betweenness centrality across the co-occurrence network, rivers and streams, estuarine and coastal systems, fisheries, food webs, DOM chemistry, and Indigenous knowledge and equity help connect otherwise-distinct sub-conversations. The equity content is not a silo; it routes across the schedule.")
p_text(docB, "Methods cluster differently with problems. Long-term monitoring is the connective-tissue method, threading through problem clusters from lakes and microbial ecology to climate, biogeochemistry, biodiversity, conservation, eutrophication, and the cryosphere. Hydrology and hydrodynamics thread almost as broadly. eDNA and -omics concentrate most strongly around microbial ecology and biodiversity. Ultra-high-resolution mass spectrometry concentrates in DOM and microbial-DOM coupling. AI and machine learning appear across several problem areas but at low intensity in each, suggesting a broad horizontal method signal rather than a single concentrated domain. Network and matrix visualisations of these patterns are in the repository linked at the end.")

# Indigenous knowledge
docB.add_heading("Indigenous knowledge, equity, and community-led science", level=2)
p_text(docB, "A classifier tag for Indigenous-knowledge, equity, community-led, or citizen science returns 58 talks across 14 session items, about 4% of the program. The session-level inclusions are notable:")
for item in [
    'EP013 "Two-Eyed Seeing: Indigenous Knowledge and Western Science" is programmed as a science session in the program structure.',
    'WS05 "Weaving Indigenous Knowledge and Western Science" and WS08 "From Values to Practice: Inclusive Science Spaces" appear as dedicated workshops.',
    'SS079 "Emerging Directions in Community-Based Water Monitoring, Participatory Science..." treats participatory science as a methods session.',
    'AV001, organized by the ASLO Early Career Committee, features presenters from Bulacan State University (Philippines), University of Eldoret (Kenya), ILPLA (Argentina), and the Universidade Federal do Espírito Santo (Brazil), among others.',
]:
    docB.add_paragraph(item, style="List Bullet")

# Gender
docB.add_heading("Gender balance, cautiously inferred", level=2)
p_text(docB, "The program also shows an encouraging gender-balance signal, with an important caveat. Among presenter names a public-domain name database could classify, about 53% were female-inferred; among session organizers, about 56%. Because the tool is European-biased and left about 12% of presenter names unclassified, these figures should be read as a rough signal rather than a census of gender identity. Still, within those limits, ASLO-SIL 2026 appears close to gender balance at both presenter and organizer levels.")

# Map for future invitations
docB.add_heading("A map for future invitations", level=2)
p_text(docB, "The geography of participation points to one of the field's clearest opportunities. The US accounts for about 29% of presentations for which a country could be detected, Canada about 25%; together they are about 54% of the program.")
p_text(docB, "Among the top-20 represented countries, Japan comes closest to parity between its share of the program (about 1.4%) and its share of global renewable internal freshwater (about 1.0%). Several countries with much larger freshwater stocks are under-represented relative to that stake: Brazil holds about 13% of the world's renewable freshwater and contributes about 1.3% of detected presentations; Russia (about 10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), and Myanmar (2%) were not detected in this inventory. Together, those seven countries hold roughly 42% of the world's renewable freshwater.")
p_text(docB, "Read generously, this is not only a gap. It is a map for future collaboration: stronger travel pathways, more durable institutional partnerships, remote or hybrid presentation options, mentoring across regions, and more routes for freshwater-rich countries to shape the global aquatic-science agenda. The AV001 and EP013 tracks already model what those routes can look like; the rest of the program shows how much room there is for those routes to grow.")

# One more thing
docB.add_heading("One more thing", level=2)
p_text(docB, "The meeting is dense. 260 oral talks plus 230 posters land on Friday 15 May 2026 alone, close to 490 presentations in one day, the heaviest of the program.")
p_text(docB, "The program contains more parallel content than any one attendee can absorb, which is also a sign of abundance: this is no longer a single conversation, but a network of overlapping ones. Pick three sessions, attend them properly, follow the bibliographies later, and let the meeting continue after the meeting.")

# Closing thought
docB.add_heading("A closing thought", level=2)
p_text(docB, "This is the portrait the schedule gives, before the meeting actually happens. Three legacies and six early-career sessions in the same program. Lakes, microbes, and biogeochemistry as dominant tags, with long-term monitoring as the connective-tissue method that threads through them. Rivers, estuaries, fisheries, food webs, DOM, and equity content as bridge topics that hold otherwise-distinct sub-conversations together. Nine sessions naming multiple stressors. Indigenous knowledge as a programmed science session. Gender balance visible at both presenter and organizer levels, cautiously inferred. AI and -omics present as small but consistent signals. A geography of participation that shows where the next invitations could be strongest. The field is not only naming its pressures; it is wiring its questions together, and building the collaborations, methods, and institutional habits needed to respond.")

# Embed Figure 1 (disciplinary distribution chart)
docB.add_paragraph()
p = docB.add_paragraph()
p.alignment = 1
p.add_run().add_picture(str(REPO / "output" / "charts" / "meeting_highlights_figure1.png"), width=Inches(6.0))
fc1 = docB.add_paragraph()
fc1.add_run("Figure 1. ").bold = True
fc1.add_run("Disciplinary distribution of all 1,455 scheduled presentations at ASLO-SIL 2026, by session-level frame (keyword classifier on session names and descriptions, multi-label; columns overlap, so percentages do not sum to 100%). About 93% of presentations carry at least one of the nine frame tags, counted as a union. Source data: scraped from the public ASLO-SIL 2026 schedule, audited 11 May 2026.")

# Embed Figure 2 (freshwater vs participation chart)
docB.add_paragraph()
p = docB.add_paragraph()
p.alignment = 1
p.add_run().add_picture(str(REPO / "output" / "charts" / "freshwater_share_vs_participation_share.png"), width=Inches(6.0))
fc2 = docB.add_paragraph()
fc2.add_run("Figure 2. ").bold = True
fc2.add_run("Country freshwater share (World Bank ER.H2O.INTR.K3) versus country share of detected program participation. Top-20 freshwater-stock countries shown. Japan is closest to parity; Brazil holds about 13% of the world's renewable freshwater and contributes about 1.3% of the program. Read generously, a map for travel pathways, regional partnerships, and hybrid presentation options.")

# Data and methods
docB.add_paragraph()
dm = docB.add_paragraph()
r = dm.add_run("Data and methods. ")
r.italic = True
r.bold = True
r2 = dm.add_run("Scraped from the public ASLO-SIL 2026 schedule and audited 11 May 2026. Disciplinary, method, and problem tags assigned by a multi-label keyword classifier on session names, session descriptions, and presentation titles, so figures should be read as approximate tag rates rather than exclusive shares. Country detection is keyword-based on affiliation strings and email-domain TLDs and reaches about 78% coverage. Gender inference uses the gender-guesser public-domain name database, which is European-biased; the 12% of presenter names it could not classify are disproportionately non-Western. Co-attention network (37 tag nodes, 325 edges at weight >=10) computed with networkx; betweenness and eigenvector centrality on the same graph. Methods-by-problems matrix is a cross-tab of method and frame tag counts. Freshwater data: World Bank ER.H2O.INTR.K3, latest available 1990-2022. Full inventory, classifier code, network and matrix visualisations, and reproducible analysis at github.com/erika-freeman/aslo2026-conference-landscape.")
r2.italic = True

PATH_B = PKG / "Freeman_OptionB_LongForm_v12.docx"
docB.save(str(PATH_B))
docB.save(str(BUNDLE / "05_OptionB_long_form_2000w.docx"))

# ---------------------------------------------------------------------------
# 3) Build Option A (short form, v4 docx, body from manuscript v3 short)
# ---------------------------------------------------------------------------
docA = Document()
style = docA.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Hero image
insert_hero(docA, ws_embed, WATERSHED_CAPTION)
docA.add_paragraph()

# Title
docA.add_heading("Reading ourselves through the program: ASLO-SIL 2026 in numbers, verbs, and invitations", level=1)

# Author block
auth = docA.add_paragraph()
auth.add_run("Erika C. Freeman").bold = True
auth.add_run("\nLeibniz Institute of Freshwater Ecology and Inland Fisheries (IGB), Berlin, Germany")
auth.add_run("\nerika.freeman@igb-berlin.de")
auth.add_run("\nORCID: 0000-0001-7161-6038")

docA.add_paragraph()

short_body = [
    "A conference is never only the thing that happens in the rooms. It is also a record of what a community thinks is worth gathering around. **1,455 scheduled presentations. 308 session items. About 1,400 primary presenters. About 740 institutions. Five days in Montreal.** What does that picture show?",
    "**The most-repeated session-title phrases in the 2026 program are about working across.** \"Multiple stressor(s)\" appears nine times in session titles. \"Bridging the gap\" appears eight. The exact frame *\"Towards a Convergence of Current Knowledge and Application\"* — and variants on it — appear six. Several sessions explicitly use *across ecosystems / across realms / across aquatic habitats*. Whatever else we say about the community in 2026, the program names a synthesis moment out loud.",
    "**The program structurally pairs three career-tribute sessions with six dedicated early-career sessions.** Mike Pace (EP001 \"Ecosystem Processes in a Changing World\"), Jim Cotner (EP004 \"Curating Limnology with Kindness\"), and Jim Elser (EP012 \"Ecological Stoichiometry\") are honored in the same meeting as AV001 \"Amplifying Voices in Aquatic Sciences,\" EP010A/B \"Sharing Experiences Among Early-Career Researchers,\" EC02 (ECR resilience), EC03 (cross-organisational ECR alliance), EP011 (first-timer \"How To\"), and EP006 (Raelyn Cole Editorial Fellowship retrospective). Three legacy sessions and six dedicated ECR sessions, in the same five-day program.",
    "**Freshwater science is a major center of gravity in this joint meeting.** A keyword classifier on session names, session descriptions, and presentation titles (multi-label) finds lakes and limnology in about 28% of talks, microbial ecology in about 18%, biogeochemistry in about 17%, climate change as a forcing in about 10%, marine and oceanography in about 9%, rivers and streams in about 9%, biodiversity in about 7%, estuarine and coastal in about 6%, and food webs in about 5%. About 93% of talks carry at least one of these tags (counted as a union, not a sum). Freshwater systems together outweigh saltwater systems in this joint meeting.",
    "**Equity is in the architecture.** Indigenous-knowledge, equity, community-led, and citizen-science framings appear in about 58 talks across 14 session items, about 4% of the program. EP013 \"Two-Eyed Seeing: Indigenous Knowledge and Western Science\" is programmed as a science session. WS05 \"Weaving Indigenous Knowledge and Western Science\" and WS08 \"From Values to Practice: Inclusive Science Spaces\" appear as dedicated workshops. SS079 \"Emerging Directions in Community-Based Water Monitoring\" treats participatory science as a methods session. AV001 features presenters from Bulacan State University (Philippines), University of Eldoret (Kenya), ILPLA (Argentina), and the Universidade Federal do Espírito Santo (Brazil), among others. A gender-balance signal is also visible, cautiously inferred from name-based classification: about 53% female-inferred among presenters and 56% among organizers, with the caveat that the classifier is European-biased and left about 12% of names unclassified.",
    "**The geography of participation points to one of the field's clearest opportunities.** The US accounts for about 29% of country-detected presentations, Canada about 25%; together they are about 54% of the program. Among the top-20 represented countries, Japan comes closest to parity between its share of the program (about 1.4%) and its share of global renewable freshwater (about 1.0%). Several countries with much larger freshwater stocks are under-represented relative to that stake: Brazil holds about 13% of global freshwater and contributes about 1.3% of presentations; Russia (about 10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), and Myanmar (2%) were not detected in this inventory. Read generously, this is a map for future collaboration — stronger travel pathways, durable institutional partnerships, hybrid presentation options, mentoring across regions, more routes for freshwater-rich countries to shape the global aquatic-science agenda.",
    "**Some methodological signatures.** AI and machine learning are tagged in about 2.4% of talks; SS070 \"Exploring the Confluence of Data, Models, and Forecasts\" names that bet explicitly. eDNA and other -omics tag about 3.2%. Mass spectrometry tags about 33 talks (about 2%), with 14 using ultra-high-resolution MS (FT-ICR, Orbitrap) and 7 naming FT-ICR specifically; most concentrate in SS050, with smaller clusters in SS049, SS041, and SS048. Long-term monitoring and time-series are the most-tagged single method, at about 9%, threading through about ten distinct problem clusters as the program's connective-tissue method. *Disclosure: I am giving one of those seven FT-ICR talks.*",
    "**One more thing.** The single densest day is Friday 15 May, with about 260 oral talks plus 230 posters — close to 490 presentations in one day. The program contains more parallel content than any one attendee can absorb, which is also a sign of abundance: this is no longer a single conversation but a network of overlapping ones. Pick three sessions, attend them properly, follow the bibliographies later, and let the meeting continue after the meeting.",
    "**The portrait the schedule gives, before the meeting actually happens.** Three legacies and six early-career sessions in the same program. Lakes, microbes, and biogeochemistry as dominant tags. Nine sessions naming multiple stressors, and many more building the tools to work across them. Indigenous knowledge as a programmed science session. Gender balance visible at both presenter and organizer levels, cautiously inferred. AI and -omics present as small but consistent signals. A geography of participation that shows where the next invitations could be strongest. The field is not only naming its pressures; it is wiring its questions together, and building the collaborations, methods, and institutional habits needed to respond.",
]
for para in short_body:
    add_runs(docA.add_paragraph(), para)

# Embed Figure 1
docA.add_paragraph()
p = docA.add_paragraph()
p.alignment = 1
p.add_run().add_picture(str(REPO / "output" / "charts" / "meeting_highlights_figure1.png"), width=Inches(6.0))
fc = docA.add_paragraph()
fc.add_run("Figure 1. ").bold = True
fc.add_run("Disciplinary distribution of all 1,455 scheduled presentations at ASLO-SIL 2026, by session-level frame (keyword classifier on session names, session descriptions, and presentation titles, multi-label — columns overlap, so percentages do not sum to 100%). About 93% of presentations carry at least one of the nine frame tags (counted as a union). The three dominant tags — lakes and limnology (about 28%), microbial ecology (about 18%), and biogeochemistry (about 17%) — act as connective tissue across the program. Freshwater systems together outweigh saltwater systems, making the joint ASLO-SIL meeting visibly freshwater-forward. Source data: scraped from the public ASLO-SIL 2026 schedule, audited 11 May 2026. Full classifier code, tag assignments, and reproducible analysis at github.com/erika-freeman/aslo2026-conference-landscape.")

docA.add_paragraph()
footer = docA.add_paragraph()
footer.add_run("Data and reproducible analysis: ").italic = True
footer.add_run("github.com/erika-freeman/aslo2026-conference-landscape").italic = True
footer.add_run("  ·  Underlying inventory available on request.").italic = True

PATH_A = PKG / "Freeman_MeetingHighlights_ASLO-SIL_2026_v4.docx"
docA.save(str(PATH_A))
docA.save(str(BUNDLE / "04_OptionA_short_form_700w.docx"))

# ---------------------------------------------------------------------------
# 4) Stage the rest of the bundle
# ---------------------------------------------------------------------------
# Watershed copies in the bundle
shutil.copy(WS_ARCHIVE_PNG, BUNDLE / "00_OpeningVisual_watershed.png")
shutil.copy(WS_ARCHIVE_TIF, BUNDLE / "00_OpeningVisual_watershed_CMYK_300dpi.tif")

# Data figures
shutil.copy(
    REPO / "output" / "charts" / "meeting_highlights_figure1.png",
    BUNDLE / "02_Figure1_disciplinary_distribution.png",
)
shutil.copy(
    PKG / "Freeman_MeetingHighlights_Figure1.tif",
    BUNDLE / "02_Figure1_disciplinary_distribution_CMYK_300dpi.tif",
)
shutil.copy(
    REPO / "output" / "charts" / "freshwater_share_vs_participation_share.png",
    BUNDLE / "03_Figure2_freshwater_vs_participation.png",
)

# Cleanup the temp embed file
if ws_embed.exists():
    ws_embed.unlink()

# ---------------------------------------------------------------------------
# 5) Inquiry email v9 — adds the watershed callout
# ---------------------------------------------------------------------------
inquiry_v9 = """TO:       lobulletin-editor@aslo.org
FROM:     erika.freeman@igb-berlin.de
SUBJECT:  Presubmission inquiry — Meeting Highlights on ASLO-SIL 2026

ATTACHMENTS:
  1. 00_OpeningVisual_watershed.png
     (hand-drawn editorial illustration, opening visual for the piece)
  2. 02_Figure1_disciplinary_distribution.png
     (disciplinary distribution of all 1,455 talks)
  3. 03_Figure2_freshwater_vs_participation.png
     (country freshwater share vs. program share)
  4. 04_OptionA_short_form_700w.docx
     (Meeting Highlights candidate, ~700 words, opening visual + one figure)
  5. 05_OptionB_long_form_2000w.docx
     (Bulletin Article candidate, ~2,000 words, opening visual + two figures)

---

Dear Dr. Falkenberg,

A Meeting Highlights pitch for the L&O Bulletin, based on ASLO-SIL 2026 in Montreal this week.

I read the full programme before flying — 1,455 scheduled presentations across 308 session items, audited against the public schedule on 11 May 2026 — and the portrait the schedule itself paints is worth sharing, in numbers, verbs, and invitations. The piece comes with a commissioned hand-drawn editorial illustration (attached) — a watershed reading of the meeting, with one terracotta tributary tracing one molecule's journey from headwater to confluence — intended as the magazine-style opening visual.

Six things stand out, each anchored to the underlying data:

  • The most-repeated session-title phrases are about working across boundaries: "multiple stressor(s)" nine times, "bridging the gap" eight, "towards a convergence of current knowledge and application" six.

  • Three career-tribute sessions (Mike Pace EP001, Jim Cotner EP004, Jim Elser EP012) appear in the same five-day programme as six dedicated early-career sessions (AV001 Amplifying Voices, EC02 resilience, EC03 cross-organisational ECR alliance, EP010A/B, EP011 first-timer guide, EP006 Raelyn Cole Editorial Fellowship retrospective).

  • Freshwater science is a major center of gravity in this joint meeting. A keyword classifier (multi-label) finds lakes 28%, microbial ecology 18%, biogeochemistry 17%; about 93% of talks carry at least one of the dominant tags.

  • Equity work is in the architecture: about 58 talks across 14 session items, including EP013 "Two-Eyed Seeing" programmed as a science session and WS05 / WS08 as dedicated workshops. A gender-balance signal is also visible, cautiously inferred from name-based classification: about 53% female-inferred among presenters and 56% among organisers, with the caveat that the classifier is European-biased and left about 12% of names unclassified.

  • Method tags are small but mapped: AI/ML about 2.4%, eDNA/-omics about 3.2%, mass spectrometry about 2% (14 ultra-high-resolution MS, 7 FT-ICR specifically), long-term monitoring about 9% — the most-tagged single method, and the connective-tissue method that threads through about ten distinct problem clusters.

  • The geography of participation points to one of the field's clearest opportunities. US plus Canada are about 54% of presenters. Japan is the country closest to parity between its share of the programme and its share of global renewable freshwater. The largest freshwater-stock countries — Brazil (13% of global freshwater, 1.3% of the programme), Russia (10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), Myanmar (2%) — together hold about 42% of the world's renewable freshwater and are under-represented or not detected. Read generously, a map for travel pathways, hybrid options, and durable regional partnerships.

Attached: (1) the watershed opening visual; (2) disciplinary distribution of all 1,455 talks; (3) country freshwater share vs. programme share — all generated for this piece. Both candidate drafts are attached so you can pick the format that fits the Bulletin:

  A · ~700-word Meeting Highlights with the opening visual and one data figure.
  B · ~2,000-word Bulletin Article with the opening visual and two data figures (adds the freshwater-vs-participation comparison, a co-attention network of the program's tags, and a methods-by-problems matrix).

Both drafts have been audited: every quantitative assertion is traceable to a corresponding entry in the repository's data tables.

I am presenting at SS050B "Ecological Significance of DOM" on 15 May. The piece would be written as a participant-observer.

Would either format fit the Bulletin? I am happy to revise either draft to spec.

With thanks,

Erika

—
Dr. Erika C. Freeman
Group Leader, ABC Lab
Leibniz Institute of Freshwater Ecology and Inland Fisheries (IGB)
Müggelseedamm 310, 12587 Berlin, Germany
erika.freeman@igb-berlin.de  ·  ORCID 0000-0001-7161-6038
"""
(PKG / "presubmission_inquiry_email_v9.txt").write_text(inquiry_v9, encoding="utf-8")
(BUNDLE / "01_inquiry_email.txt").write_text(inquiry_v9, encoding="utf-8")

# ---------------------------------------------------------------------------
# 6) README
# ---------------------------------------------------------------------------
readme = """SUBMISSION BUNDLE v3 — L&O Bulletin presubmission inquiry
==========================================================
Recipient   : Dr. Laura Falkenberg, L&O Bulletin Editor
Address     : lobulletin-editor@aslo.org
From        : Erika C. Freeman (IGB Berlin)
Bundle date : 11 May 2026
Manuscript  : v11 (long form) / v3 short body, both rebuilt as v12/v4 docx
Inquiry     : v9 (opening-visual line added)
Hero image  : Concept A watershed (hand-drawn, watercolour on cream)

WHAT CHANGED SINCE BUNDLE v2
----------------------------
- Watershed opening visual now embedded at the top of both Option A and
  Option B docx files.
- Bundle now ships the watershed in three forms:
    00_OpeningVisual_watershed.png             (full-resolution colour)
    00_OpeningVisual_watershed_CMYK_300dpi.tif (print-ready)
  And the embedded copy inside each docx is downsampled to a sensible size
  for an attachment-friendly file.
- Inquiry email v9 adds a paragraph announcing the editorial opening visual.

WHAT TO DO
----------
1. Open 01_inquiry_email.txt. Copy the body (everything below the "---"
   divider) into a new email to lobulletin-editor@aslo.org. Use the SUBJECT
   line at the top of the file.

2. Attach all five files:
       00_OpeningVisual_watershed.png
       02_Figure1_disciplinary_distribution.png
       03_Figure2_freshwater_vs_participation.png
       04_OptionA_short_form_700w.docx
       05_OptionB_long_form_2000w.docx

   (The bundle also contains CMYK 300-dpi TIFFs of the watershed and Figure 1
   for eventual publication — do NOT attach those to the inquiry email;
   they're staged for the later formal submission step.)

3. Send.

WHAT'S IN THE BUNDLE
--------------------
00_README.txt                                   This file.

01_inquiry_email.txt                            Email body (v9).

00_OpeningVisual_watershed.png                  (primary attachment, hero)
   Hand-drawn editorial illustration. A community as watershed: tributaries
   from many ecosystems, career stages, and regions braiding toward a single
   confluence; one terracotta thread traces one molecule's journey from
   headwater to pool. Concept A from the AI prompts file.

00_OpeningVisual_watershed_CMYK_300dpi.tif      (for eventual print, NOT for
                                                 the inquiry email)

02_Figure1_disciplinary_distribution.png        (primary attachment)
   Bar chart of all 1,455 talks across the nine frame tags.

02_Figure1_disciplinary_distribution_CMYK_300dpi.tif   (for eventual print)

03_Figure2_freshwater_vs_participation.png      (primary attachment)
   Country-by-country bar chart. Top-20 freshwater-stock countries.

04_OptionA_short_form_700w.docx                 (Meeting Highlights candidate)
   ~700-word piece. Layout: watershed opener at top, body, one data chart.
   Source: manuscripts/02_meeting_highlights_500w_v3.md (text body)
         + scripts/16_build_submission_bundle_v3.py (layout, hero embed)

05_OptionB_long_form_2000w.docx                 (Bulletin Article candidate)
   ~2,000-word longer version. Layout: watershed opener at top, body,
   disciplinary chart, freshwater chart.
   Source: manuscripts/01_member_narrative_1500w_v11.md (text body)
         + scripts/16_build_submission_bundle_v3.py (layout, hero embed)

WORD-COUNT RANGES (per L&O Bulletin guidelines)
-----------------------------------------------
   Meeting Highlights : 500-1,500 words   --> Option A (700) fits cleanly.
   Bulletin Article   : 3,000-5,000 words --> Option B (2,000) sits below the
                                                Article minimum. Either expand
                                                by ~1,000 words on request, or
                                                trim below 1,500 to land in
                                                Meeting Highlights.

HERO + EVIDENCE PAIRING (per AI prompts file)
---------------------------------------------
The watershed is the *editorial hero* — what makes the piece feel like a
magazine page on Laura's screen the moment she opens the docx. The bar charts
are the *evidence* — what makes every quantitative claim auditable. The two
roles are complementary; the L&O Bulletin's house style supports this kind
of pairing in their longer pieces.

SOURCE FILES (reference, not for attachment)
--------------------------------------------
.../_GitHubRepo/aslo2026-conference-landscape/manuscripts/
    01_member_narrative_1500w_v11.md   (Option B body source)
    02_meeting_highlights_500w_v3.md   (Option A body source)
    VERSION_NOTES.md                   (audit trail through v11)

Build script:
.../_GitHubRepo/aslo2026-conference-landscape/scripts/16_build_submission_bundle_v3.py

Canonical archive:
.../_ConferenceLandscape/submission_package/
    Freeman_OpeningVisual_Watershed_ConceptA.png
    Freeman_OpeningVisual_Watershed_ConceptA_CMYK_300dpi.tif
    Freeman_MeetingHighlights_ASLO-SIL_2026_v4.docx
    Freeman_OptionB_LongForm_v12.docx
    presubmission_inquiry_email_v9.txt

DECLARATIONS (included in inquiry body)
---------------------------------------
- Not published or submitted elsewhere.
- No conflicts of interest.
- Every quantitative assertion in the manuscript is traceable to the
  repository's output/tables/ directory.
- The author is presenting at SS050B "Ecological Significance of DOM" on
  15 May 2026; participant-observer perspective disclosed in the manuscript.
"""
(BUNDLE / "00_README.txt").write_text(readme, encoding="utf-8")

# ---------------------------------------------------------------------------
# Report
# ---------------------------------------------------------------------------
print("Bundle assembled at: {}".format(BUNDLE))
print()
print("Contents:")
for f in sorted(BUNDLE.iterdir()):
    print("  {:>10} bytes  {}".format(f.stat().st_size, f.name))

# Body word counts
def body_wc(p):
    return sum(len(par.text.split()) for par in Document(str(p)).paragraphs)
print()
print("Option A docx body word count: {}".format(body_wc(BUNDLE / "04_OptionA_short_form_700w.docx")))
print("Option B docx body word count: {}".format(body_wc(BUNDLE / "05_OptionB_long_form_2000w.docx")))

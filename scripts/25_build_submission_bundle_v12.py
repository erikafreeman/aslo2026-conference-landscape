"""
Build submission_bundle_v12/ — viewer-safe margins and image sizing.

Problem in bundle v11: the watershed hero image was 6.0 inches wide and
the page content width was also 6.0 inches (python-docx default margins
of 1.25 in left/right). The image touched the right margin exactly,
which Microsoft Word renders fine but several .docx previewers
(Google Drive's web preview in particular) misrender — sometimes
omitting images and clipping content at page boundaries.

Fix:
  - Explicit 1.0 in margins on all four sides (content width = 6.5 in).
  - Watershed embedded at 5.5 in wide (0.5 in margin breathing room).
  - Charts embedded at 5.0 in wide.
  - Watershed downsampled to 1400 px wide before embedding to keep
    the .docx file size sensible.

If after this it still looks wrong in a previewer, the fix is to
download the file from Drive and open it in Microsoft Word or
LibreOffice; web previewers continue to be unreliable for any
python-docx file with embedded images.

Outputs:
  submission_bundle_v12/
  submission_package/Freeman_ASLO-SIL_2026_Programme_Analysis_v3.docx
  submission_package/presubmission_inquiry_email_v17.txt
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from PIL import Image
import shutil
import re

REPO = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_GitHubRepo\aslo2026-conference-landscape")
CONFLAND = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_ConferenceLandscape")
BUNDLE = CONFLAND / "submission_bundle_v12"
PKG = CONFLAND / "submission_package"
BUNDLE.mkdir(parents=True, exist_ok=True)

WS_ARCHIVE_PNG = PKG / "Freeman_OpeningVisual_Watershed_ConceptA.png"
WS_ARCHIVE_TIF = PKG / "Freeman_OpeningVisual_Watershed_ConceptA_CMYK_300dpi.tif"
if not WS_ARCHIVE_PNG.exists():
    raise SystemExit("Watershed PNG missing.")

# Downsample watershed to 1400 px wide for embedding (smaller file size,
# still high enough resolution for ~5.5 inches at 250 dpi).
ws_embed = BUNDLE / "_tmp_watershed_for_embed.png"
img_rgb = Image.open(WS_ARCHIVE_PNG).convert("RGB")
if img_rgb.width > 1400:
    ratio = 1400 / img_rgb.width
    img_rgb = img_rgb.resize((1400, int(img_rgb.height * ratio)), Image.LANCZOS)
img_rgb.save(ws_embed, format="PNG", optimize=True)

# ---------------------------------------------------------------------------
# Manuscript .docx formatting defaults + viewer-safe margins
# ---------------------------------------------------------------------------
def apply_manuscript_format(doc, body_font="Times New Roman", body_size_pt=12):
    style = doc.styles["Normal"]
    style.font.name = body_font
    style.font.size = Pt(body_size_pt)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    rPr = style.element.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(k), body_font)
    rPr.append(rFonts)

    for hname, size in [("Heading 1", 16), ("Heading 2", 13)]:
        try:
            hstyle = doc.styles[hname]
            hstyle.font.name = body_font
            hstyle.font.size = Pt(size)
            hstyle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            hrPr = hstyle.element.get_or_add_rPr()
            existing = hrPr.find(qn("w:rFonts"))
            if existing is not None:
                hrPr.remove(existing)
            hr = OxmlElement("w:rFonts")
            for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                hr.set(qn(k), body_font)
            hrPr.append(hr)
        except KeyError:
            pass

    for lname in ("List Bullet", "List Number"):
        try:
            ls = doc.styles[lname]
            ls.font.name = body_font
            ls.font.size = Pt(body_size_pt)
            ls.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        except KeyError:
            pass

    # Override default margins to 1.0 inch on all four sides
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    # Letter (default) — explicit for clarity
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)

    sectPr = section._sectPr
    existing = sectPr.find(qn("w:lnNumType"))
    if existing is not None:
        sectPr.remove(existing)
    lnNumType = OxmlElement("w:lnNumType")
    lnNumType.set(qn("w:countBy"), "1")
    lnNumType.set(qn("w:start"), "1")
    lnNumType.set(qn("w:restart"), "continuous")
    lnNumType.set(qn("w:distance"), "360")  # 0.25 in in twips
    sectPr.append(lnNumType)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def add_runs(paragraph, text):
    pat = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        else:
            r = paragraph.add_run(token[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def p(doc, text, style=None):
    par = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_runs(par, text)
    return par

WATERSHED_CAPTION = (
    "Opening visual: A community as watershed. Tributaries from ecosystems, "
    "career stages, and regions braid toward a shared confluence; a "
    "terracotta thread traces one molecule's journey from headwater to pool."
)

# Image widths: well inside the 6.5 in content area
HERO_WIDTH = Inches(5.5)
CHART_WIDTH = Inches(5.0)

def insert_hero(doc, image_path, caption):
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run().add_picture(str(image_path), width=HERO_WIDTH)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(10)

def embed_chart(doc, image_path, caption_label, caption_body):
    doc.add_paragraph()
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run().add_picture(str(image_path), width=CHART_WIDTH)
    cap = doc.add_paragraph()
    cap.add_run(caption_label).bold = True
    cap.add_run(caption_body)

# ===========================================================================
# Build manuscript
# ===========================================================================
doc = Document()
apply_manuscript_format(doc)

insert_hero(doc, ws_embed, WATERSHED_CAPTION)
doc.add_paragraph()

doc.add_heading("ASLO-SIL 2026 in Numbers: Trends, Gaps, and Takeaways from the Programme", level=1)

auth = doc.add_paragraph()
auth.add_run("Erika C. Freeman").bold = True
auth.add_run("\nLeibniz Institute of Freshwater Ecology and Inland Fisheries (IGB), Berlin, Germany")
auth.add_run("\nerika.freeman@igb-berlin.de  ·  ORCID 0000-0001-7161-6038")

doc.add_paragraph()

p(doc, "1,455 scheduled presentations. 308 session items. Approximately 1,400 primary presenters across 740 institutions. Five days in Montreal.")
p(doc, "A conference programme is, in effect, a record of collective focus. Reading the session titles, keywords, and method tags for ASLO-SIL 2026 closely gives a sense of what the field is currently spending its attention on. The schedule turns up some surprising trends, some areas of progress, and some structural gaps worth noticing.")
p(doc, "Here is what the data suggests about where the field is, and some of the directions it appears to be moving in.")

doc.add_heading("Unexpected Trends", level=2)
p(doc, "**Long-term monitoring outpaces AI.** The most-tagged single method in the programme is not artificial intelligence; it is long-term monitoring. Time-series and long-term monitoring feature in about 9% of the talks, threading through distinct problem clusters from lakes and microbial ecology to climate, biogeochemistry, and the cryosphere. AI and machine learning are present (about 2.4%) but spread thin. The story here is not that \"AI is changing the field,\" but that patient, multi-decade observation is the foundation of much of the field's work, with AI showing up as a supplementary tool.")
p(doc, "**Equity is integrated, not siloed.** Framings around Indigenous knowledge, equity, community-led initiatives, and citizen science appear in 58 talks across 14 session items (about 4% of the programme). The percentage is small, but the structural role in the programme's tag network is larger than the share suggests. Indigenous knowledge and equity sit alongside rivers, estuaries, fisheries, food webs, and DOM chemistry as \"bridge topics\" connecting otherwise distinct sub-conversations. Sessions like EP013 (\"Two-Eyed Seeing\") are programmed as core science, not adjuncts. Equity content isn't happening on the margins; it is routing directly through the main programme.")
p(doc, "**Lakes are a key convergence point.** Lakes are not just the largest category. They are also one of the main confluences where microbial ecology, climate change, long-term monitoring, and biogeochemistry repeatedly intersect. In a joint ASLO and SIL meeting it would be easy to read this freshwater dominance as \"limnology won this round.\" It reads more usefully as: lakes are one of the systems through which many of the field's broadest questions are currently being asked.")

doc.add_heading("Emerging Themes", level=2)
p(doc, "**An emphasis on synthesis.** The most repeated phrases in session titles are about crossing boundaries. \"Multiple stressor(s)\" appears nine times, \"Bridging the gap\" eight times, and \"Convergence\" six times. The programme has a visible push toward synthesis across ecosystems and habitats.")
p(doc, "**Legacy and the early-career cohort, side by side.** Three career-tribute sessions for senior scientists (Mike Pace, Jim Cotner, and Jim Elser) share the schedule with six dedicated early-career (ECR) sessions. Whether intentional or not, the programme structurally pairs the establishment with the next generation.")
p(doc, "**Microbial, biogeochemistry, and DOM integration.** Sessions like SS058 (\"Microbial-DOM Coupling\") and SS050 (\"Ecological Significance of DOM\") treat reciprocal microbial-molecular dynamics as a coherent research front.")
p(doc, "**Forecasting and monitoring in the same conversation.** SS070 puts long-term observation, mechanistic modelling, and forward prediction in the same room. Prediction appears to be moving inside the science itself, rather than sitting as a translation layer applied after the fact.")

doc.add_heading("Areas of Progress", level=2)
p(doc, "The programme quietly shows a strong gender-balance signal: approximately 53% female-inferred among classifiable presenter names and 56% among organisers (using a European-biased classifier with a 12% unclassified rate). Within those limits, ASLO-SIL 2026 looks close to parity, a benchmark still uncommon across STEM.")
p(doc, "Roughly 93% of presentations carry at least one of the nine dominant frame tags. Despite the volume of talks, the field shows breadth and cohesion rather than fragmentation.")

doc.add_heading("Programme Density and the Challenge of Scale", level=2)
p(doc, "The heaviest day of the meeting is Friday, 15 May: roughly 260 oral talks and 230 posters land within 24 hours. That volume means schedule conflicts and missed sessions for almost everyone.")
p(doc, "The density reads more naturally as a feature of a growing field than as a checklist to clear. Picking a few sessions to attend properly, and leaving room for the rest of the conversation to continue afterwards, is one way through.")

doc.add_heading("Geographic Gaps and Opportunities", level=2)
p(doc, "The geography of participation points to one of the clearer opportunities the programme surfaces. The US accounts for about 29% of presentations for which a country could be detected, and Canada about 25%. Together, they make up 54% of the programme.")
p(doc, "When share-of-programme is set next to share of global renewable internal freshwater stocks, Japan comes closest to parity (1.4% of the programme vs. 1.0% of global freshwater). Countries with substantial freshwater resources are notably under-represented. Brazil holds roughly 13% of the world's renewable freshwater and contributes 1.3% of detected presentations. Russia (10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), and Myanmar (2%) were not detected in this inventory. Together, those seven countries hold roughly 42% of the world's renewable freshwater.")
p(doc, "Read generously, this is a map for future collaboration. The programmatic levers (travel pathways, hybrid options, regional hubs) sit alongside deeper ones (how international authorship is valued, how funding agencies underwrite first-author opportunities, and how the published record represents work in languages and venues beyond the dominant ones). The two scales of lever tend to move together.")

doc.add_heading("Navigating the Meeting", level=2)
p(doc, "For attendees looking for practical routes through the schedule:")
for item in [
    "**Microbial-DOM-biogeochemistry:** SS050, SS058, SS002 (gas fluxes), and the isotope-tracing work in SS041 form a coherent track. Most mass-spectrometry talks cluster here.",
    "**Long-term monitoring:** Follow the method tag in the abstracts rather than session titles; it threads across lakes, climate, biodiversity, and the cryosphere.",
    "**Equity and Indigenous knowledge:** EP013 (Two-Eyed Seeing), WS05, WS08, SS079, and AV001 form a coherent, week-long arc.",
    "**Scale-bridging:** The nine \"multiple stressors\" and \"bridging the gap\" sessions are the map, anchored by SS082 (mesocosms) and SS070 (modelling/forecasts).",
    "**Cold and dry systems:** SS011, SS013, SS015, and SS019 function as a single conversation about climate-driven change, split across four sessions.",
]:
    p(doc, item, style="List Bullet")

doc.add_heading("Notes for Future Programmes", level=2)
p(doc, "A few signals from the data that future programmes might find useful, offered with no insider knowledge of why the current schedule looks the way it does:")
for item in [
    "**ECR infrastructure.** Six dedicated early-career sessions in five days is real structural support, and worth carrying forward.",
    "**Adjacent scheduling.** Tribute sessions scheduled adjacent to early-career tracks would let attendees walk straight from celebrating the past into the conversations that extend it.",
    "**Pacing.** The Friday-density signal could be eased over time by peak-day spacing, longer breaks between parallel blocks, or hybrid options for clash relief.",
    "**Bridge-session visibility.** Cross-disciplinary sessions like EP013 and AV001 carry a lot in the tag network for their size. Cross-track signposting would help them be easier to find.",
]:
    p(doc, item, style="List Bullet")

doc.add_heading("Final Thoughts", level=2)
p(doc, "The data points to a community crossing disciplinary boundaries, scaling its methods, pairing legacy with youth, and building equity into its scientific core. Have a brilliant time in Montreal.")

embed_chart(
    doc,
    REPO / "output" / "charts" / "meeting_highlights_figure1.png",
    "Figure 1. ",
    "Disciplinary distribution of all 1,455 scheduled presentations at ASLO-SIL 2026, by session-level frame. (Keyword classifier on session names, session descriptions, and presentation titles; multi-label. Columns overlap; percentages do not sum to 100%.) Approximately 93% of presentations carry at least one of the nine frame tags. Source: scraped from the public ASLO-SIL 2026 schedule, audited 11 May 2026.",
)
embed_chart(
    doc,
    REPO / "output" / "charts" / "freshwater_share_vs_participation_share.png",
    "Figure 2. ",
    "Country freshwater share (World Bank ER.H2O.INTR.K3) versus country share of presentations for which a country could be detected. Top-20 freshwater-stock countries shown. Japan is closest to parity; Brazil holds about 13% of the world's renewable freshwater and contributes about 1.3% of the programme.",
)

doc.add_paragraph()
dm = doc.add_paragraph()
r1 = dm.add_run("Data and methods: "); r1.italic = True; r1.bold = True
r2 = dm.add_run("Scraped from the public ASLO-SIL 2026 schedule and audited 11 May 2026. Disciplinary, method, and problem tags assigned by a multi-label keyword classifier. Country detection reaches approx. 78% coverage. Gender inference utilises the European-biased gender-guesser database (12% unclassified). Full inventory, classifier code, network/matrix visualisations, and reproducible analysis available at github.com/erika-freeman/aslo2026-conference-landscape.")
r2.italic = True

byline = doc.add_paragraph()
br = byline.add_run("Erika C. Freeman · IGB Berlin · SS050B \"Ecological Significance of Dissolved Organic Matter\" · 15 May 2026 · Palais des congrès de Montréal.")
br.italic = True

DOCX_NAME = "Freeman_ASLO-SIL_2026_Programme_Analysis_v3.docx"
PATH = PKG / DOCX_NAME
doc.save(str(PATH))
doc.save(str(BUNDLE / "04_Manuscript_Freeman_1100w.docx"))

# ===========================================================================
# Stage figures and watershed
# ===========================================================================
shutil.copy(WS_ARCHIVE_PNG, BUNDLE / "00_OpeningVisual_watershed.png")
shutil.copy(WS_ARCHIVE_TIF, BUNDLE / "00_OpeningVisual_watershed_CMYK_300dpi.tif")
shutil.copy(REPO / "output" / "charts" / "meeting_highlights_figure1.png",
            BUNDLE / "02_Figure1_disciplinary_distribution.png")
shutil.copy(PKG / "Freeman_MeetingHighlights_Figure1.tif",
            BUNDLE / "02_Figure1_disciplinary_distribution_CMYK_300dpi.tif")
shutil.copy(REPO / "output" / "charts" / "freshwater_share_vs_participation_share.png",
            BUNDLE / "03_Figure2_freshwater_vs_participation.png")

if ws_embed.exists():
    ws_embed.unlink()

# ===========================================================================
# Inquiry email (same v16 content, no changes)
# ===========================================================================
shutil.copy(PKG / "presubmission_inquiry_email_v16.txt", BUNDLE / "01_inquiry_email.txt")

# Make a canonical v17 copy so the inquiry version stays in sync with bundle v12
shutil.copy(PKG / "presubmission_inquiry_email_v16.txt", PKG / "presubmission_inquiry_email_v17.txt")

# ===========================================================================
# README
# ===========================================================================
readme = """SUBMISSION BUNDLE v12  |  L&O Bulletin presubmission inquiry
=============================================================
Recipient   : Dr. Laura Falkenberg, L&O Bulletin Editor
Address     : lobulletin-editor@aslo.org
From        : Erika C. Freeman (IGB Berlin)
Bundle date : 12 May 2026
Manuscript  : v18 (~1,100 words)
Inquiry     : v16 (unchanged from v11)

WHAT CHANGED SINCE BUNDLE v11
-----------------------------
Layout fix. Bundle v11 used 1.25-inch left/right margins (python-docx
default) with a 6.0-inch wide hero image. The image sat exactly at the
right margin, which Microsoft Word rendered correctly but some web-based
.docx previewers (including Google Drive's) misrendered, dropping images
and clipping content at page breaks.

In bundle v12:
  - Margins are explicitly 1.0 inch on all four sides (content area 6.5
    inches wide).
  - Watershed image embedded at 5.5 inches wide (0.5 in of breathing
    room from the right margin).
  - Chart figures embedded at 5.0 inches wide (0.75 in of breathing
    room).
  - Watershed downsampled to 1400 px wide before embedding (smaller
    .docx file, same display quality at 250 dpi).

If a previewer still misrenders the file, download it from Drive and
open in Microsoft Word or LibreOffice. Web-based .docx previewers
remain unreliable for any file with embedded images.

WHAT TO DO
----------
1. Open 01_inquiry_email.txt. Copy the body (below the "---" divider) into
   a new email to lobulletin-editor@aslo.org. Use the SUBJECT line at top.

2. Attach all four files:
       00_OpeningVisual_watershed.png
       02_Figure1_disciplinary_distribution.png
       03_Figure2_freshwater_vs_participation.png
       04_Manuscript_Freeman_1100w.docx

3. Send.

WHAT'S IN THE BUNDLE
--------------------
00_README.txt                                   This file.
01_inquiry_email.txt                            Email body.
00_OpeningVisual_watershed.png                  Opening visual.
00_OpeningVisual_watershed_CMYK_300dpi.tif      Print-ready.
02_Figure1_disciplinary_distribution.png        Data figure 1.
02_Figure1_disciplinary_distribution_CMYK_300dpi.tif
03_Figure2_freshwater_vs_participation.png      Data figure 2.
04_Manuscript_Freeman_1100w.docx                Full draft.

Canonical archive (submission_package/):
    Freeman_ASLO-SIL_2026_Programme_Analysis_v3.docx
    presubmission_inquiry_email_v17.txt

Build script:
.../_GitHubRepo/aslo2026-conference-landscape/scripts/25_build_submission_bundle_v12.py
"""
(BUNDLE / "00_README.txt").write_text(readme, encoding="utf-8")

# ===========================================================================
# Diagnostic verification
# ===========================================================================
import zipfile
print("=== Verification: bundle v12 manuscript ===")
docx_path = BUNDLE / "04_Manuscript_Freeman_1100w.docx"
d = Document(str(docx_path))
section = d.sections[0]
print("Page: {:.2f} x {:.2f} in".format(section.page_width.inches, section.page_height.inches))
print("Margins: top={:.2f} bottom={:.2f} left={:.2f} right={:.2f} in".format(
    section.top_margin.inches, section.bottom_margin.inches,
    section.left_margin.inches, section.right_margin.inches))
content_w = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
print("Content width: {:.2f} in".format(content_w))
print()
print("Inline images:")
for i, sh in enumerate(d.inline_shapes, 1):
    print("  {}: {:.2f} x {:.2f} in (margin from right edge: {:.2f} in)".format(
        i, sh.width.inches, sh.height.inches, content_w - sh.width.inches))
print()
with zipfile.ZipFile(str(docx_path)) as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    print("Embedded media in .docx zip:")
    for m in media:
        info = z.getinfo(m)
        print("  - {} ({:,} bytes)".format(m, info.file_size))
print()
print("File size: {:,} bytes ({:.2f} MB)".format(docx_path.stat().st_size, docx_path.stat().st_size / 1024 / 1024))

print()
print("Bundle contents:")
for f in sorted(BUNDLE.iterdir()):
    print("  {:>10,} bytes  {}".format(f.stat().st_size, f.name))

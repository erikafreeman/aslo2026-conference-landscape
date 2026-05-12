"""
Build submission_bundle_v11/ — tone pass.

Manuscript v17 -> v18. Same structure, softer voice:
  - Superlatives reduced ("ultimate," "vastly," "massive," "remarkably,"
    "completely," "impressive," "only," "merely") removed or softened.
  - Imperative "attendees should…" reframed as observation.
  - Recommendations section retitled "Notes for Future Programmes" with
    an explicit caveat about no insider knowledge of the schedule. Each
    bullet recast from imperative to observation.
  - "Final Thoughts" depuffed (no "actively," "working to").
  - Watershed image still carries no provenance claim.

Outputs:
  submission_bundle_v11/
  submission_package/Freeman_ASLO-SIL_2026_Programme_Analysis_v2.docx
  submission_package/presubmission_inquiry_email_v16.txt
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from PIL import Image
import shutil
import re

REPO = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_GitHubRepo\aslo2026-conference-landscape")
CONFLAND = Path(r"G:\My Drive\3-WorkAndPurpose\0-ABC_Lab\06-Teaching_Outreach\1-ConferenceTalks\ASLO2026\_ConferenceLandscape")
BUNDLE = CONFLAND / "submission_bundle_v11"
PKG = CONFLAND / "submission_package"
BUNDLE.mkdir(parents=True, exist_ok=True)

WS_ARCHIVE_PNG = PKG / "Freeman_OpeningVisual_Watershed_ConceptA.png"
WS_ARCHIVE_TIF = PKG / "Freeman_OpeningVisual_Watershed_ConceptA_CMYK_300dpi.tif"
if not WS_ARCHIVE_PNG.exists():
    raise SystemExit("Watershed PNG missing.")

ws_embed = BUNDLE / "_tmp_watershed_for_embed.png"
img_rgb = Image.open(WS_ARCHIVE_PNG).convert("RGB")
if img_rgb.width > 1600:
    ratio = 1600 / img_rgb.width
    img_rgb = img_rgb.resize((1600, int(img_rgb.height * ratio)), Image.LANCZOS)
img_rgb.save(ws_embed, format="PNG", optimize=True)

# ---------------------------------------------------------------------------
# Manuscript .docx formatting defaults
# ---------------------------------------------------------------------------
def apply_manuscript_format(doc, body_font="Times New Roman", body_size_pt=12):
    style = doc.styles["Normal"]
    style.font.name = body_font
    style.font.size = Pt(body_size_pt)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    rPr = style.element.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(k), body_font)
    rPr.append(rFonts)

    for hname, size in [("Heading 1", 16), ("Heading 2", 13)]:
        try:
            hstyle = doc.styles[hname]
            hstyle.font.name = body_font
            hstyle.font.size = Pt(size)
            hstyle.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            hrPr = hstyle.element.get_or_add_rPr()
            existing = hrPr.find(qn("w:rFonts"))
            if existing is not None:
                hrPr.remove(existing)
            hr = OxmlElement("w:rFonts")
            for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                hr.set(qn(k), body_font)
            hrPr.append(hr)
        except KeyError:
            pass

    for lname in ("List Bullet", "List Number"):
        try:
            ls = doc.styles[lname]
            ls.font.name = body_font
            ls.font.size = Pt(body_size_pt)
            ls.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        except KeyError:
            pass

    section = doc.sections[0]
    sectPr = section._sectPr
    existing = sectPr.find(qn("w:lnNumType"))
    if existing is not None:
        sectPr.remove(existing)
    lnNumType = OxmlElement("w:lnNumType")
    lnNumType.set(qn("w:countBy"), "1")
    lnNumType.set(qn("w:start"), "1")
    lnNumType.set(qn("w:restart"), "continuous")
    lnNumType.set(qn("w:distance"), "360")
    sectPr.append(lnNumType)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def add_runs(paragraph, text):
    pat = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        else:
            r = paragraph.add_run(token[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def p(doc, text, style=None):
    par = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_runs(par, text)
    return par

WATERSHED_CAPTION = (
    "Opening visual: A community as watershed. Tributaries from ecosystems, "
    "career stages, and regions braid toward a shared confluence; a "
    "terracotta thread traces one molecule's journey from headwater to pool."
)

def insert_hero(doc, image_path, caption):
    pp = doc.add_paragraph(); pp.alignment = 1
    pp.add_run().add_picture(str(image_path), width=Inches(6.0))
    cap = doc.add_paragraph(); cap.alignment = 1
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(10)

def embed_chart(doc, image_path, caption_label, caption_body):
    doc.add_paragraph()
    pp = doc.add_paragraph(); pp.alignment = 1
    pp.add_run().add_picture(str(image_path), width=Inches(5.5))
    cap = doc.add_paragraph()
    cap.add_run(caption_label).bold = True
    cap.add_run(caption_body)

# ===========================================================================
# Build manuscript
# ===========================================================================
doc = Document()
apply_manuscript_format(doc)

insert_hero(doc, ws_embed, WATERSHED_CAPTION)
doc.add_paragraph()

doc.add_heading("ASLO-SIL 2026 in Numbers: Trends, Gaps, and Takeaways from the Programme", level=1)

auth = doc.add_paragraph()
auth.add_run("Erika C. Freeman").bold = True
auth.add_run("\nLeibniz Institute of Freshwater Ecology and Inland Fisheries (IGB), Berlin, Germany")
auth.add_run("\nerika.freeman@igb-berlin.de  ·  ORCID 0000-0001-7161-6038")

doc.add_paragraph()

p(doc, "1,455 scheduled presentations. 308 session items. Approximately 1,400 primary presenters across 740 institutions. Five days in Montreal.")
p(doc, "A conference programme is, in effect, a record of collective focus. Reading the session titles, keywords, and method tags for ASLO-SIL 2026 closely gives a sense of what the field is currently spending its attention on. The schedule turns up some surprising trends, some areas of progress, and some structural gaps worth noticing.")
p(doc, "Here is what the data suggests about where the field is, and some of the directions it appears to be moving in.")

doc.add_heading("Unexpected Trends", level=2)
p(doc, "**Long-term monitoring outpaces AI.** The most-tagged single method in the programme is not artificial intelligence; it is long-term monitoring. Time-series and long-term monitoring feature in about 9% of the talks, threading through distinct problem clusters from lakes and microbial ecology to climate, biogeochemistry, and the cryosphere. AI and machine learning are present (about 2.4%) but spread thin. The story here is not that \"AI is changing the field,\" but that patient, multi-decade observation is the foundation of much of the field's work, with AI showing up as a supplementary tool.")
p(doc, "**Equity is integrated, not siloed.** Framings around Indigenous knowledge, equity, community-led initiatives, and citizen science appear in 58 talks across 14 session items (about 4% of the programme). The percentage is small, but the structural role in the programme's tag network is larger than the share suggests. Indigenous knowledge and equity sit alongside rivers, estuaries, fisheries, food webs, and DOM chemistry as \"bridge topics\" connecting otherwise distinct sub-conversations. Sessions like EP013 (\"Two-Eyed Seeing\") are programmed as core science, not adjuncts. Equity content isn't happening on the margins; it is routing directly through the main programme.")
p(doc, "**Lakes are a key convergence point.** Lakes are not just the largest category. They are also one of the main confluences where microbial ecology, climate change, long-term monitoring, and biogeochemistry repeatedly intersect. In a joint ASLO and SIL meeting it would be easy to read this freshwater dominance as \"limnology won this round.\" It reads more usefully as: lakes are one of the systems through which many of the field's broadest questions are currently being asked.")

doc.add_heading("Emerging Themes", level=2)
p(doc, "**An emphasis on synthesis.** The most repeated phrases in session titles are about crossing boundaries. \"Multiple stressor(s)\" appears nine times, \"Bridging the gap\" eight times, and \"Convergence\" six times. The programme has a visible push toward synthesis across ecosystems and habitats.")
p(doc, "**Legacy and the early-career cohort, side by side.** Three career-tribute sessions for senior scientists (Mike Pace, Jim Cotner, and Jim Elser) share the schedule with six dedicated early-career (ECR) sessions. Whether intentional or not, the programme structurally pairs the establishment with the next generation.")
p(doc, "**Microbial, biogeochemistry, and DOM integration.** Sessions like SS058 (\"Microbial-DOM Coupling\") and SS050 (\"Ecological Significance of DOM\") treat reciprocal microbial-molecular dynamics as a coherent research front.")
p(doc, "**Forecasting and monitoring in the same conversation.** SS070 puts long-term observation, mechanistic modelling, and forward prediction in the same room. Prediction appears to be moving inside the science itself, rather than sitting as a translation layer applied after the fact.")

doc.add_heading("Areas of Progress", level=2)
p(doc, "The programme quietly shows a strong gender-balance signal: approximately 53% female-inferred among classifiable presenter names and 56% among organisers (using a European-biased classifier with a 12% unclassified rate). Within those limits, ASLO-SIL 2026 looks close to parity, a benchmark still uncommon across STEM.")
p(doc, "Roughly 93% of presentations carry at least one of the nine dominant frame tags. Despite the volume of talks, the field shows breadth and cohesion rather than fragmentation.")

doc.add_heading("Programme Density and the Challenge of Scale", level=2)
p(doc, "The heaviest day of the meeting is Friday, 15 May: roughly 260 oral talks and 230 posters land within 24 hours. That volume means schedule conflicts and missed sessions for almost everyone.")
p(doc, "The density reads more naturally as a feature of a growing field than as a checklist to clear. Picking a few sessions to attend properly, and leaving room for the rest of the conversation to continue afterwards, is one way through.")

doc.add_heading("Geographic Gaps and Opportunities", level=2)
p(doc, "The geography of participation points to one of the clearer opportunities the programme surfaces. The US accounts for about 29% of presentations for which a country could be detected, and Canada about 25%. Together, they make up 54% of the programme.")
p(doc, "When share-of-programme is set next to share of global renewable internal freshwater stocks, Japan comes closest to parity (1.4% of the programme vs. 1.0% of global freshwater). Countries with substantial freshwater resources are notably under-represented. Brazil holds roughly 13% of the world's renewable freshwater and contributes 1.3% of detected presentations. Russia (10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), and Myanmar (2%) were not detected in this inventory. Together, those seven countries hold roughly 42% of the world's renewable freshwater.")
p(doc, "Read generously, this is a map for future collaboration. The programmatic levers (travel pathways, hybrid options, regional hubs) sit alongside deeper ones (how international authorship is valued, how funding agencies underwrite first-author opportunities, and how the published record represents work in languages and venues beyond the dominant ones). The two scales of lever tend to move together.")

doc.add_heading("Navigating the Meeting", level=2)
p(doc, "For attendees looking for practical routes through the schedule:")
for item in [
    "**Microbial-DOM-biogeochemistry:** SS050, SS058, SS002 (gas fluxes), and the isotope-tracing work in SS041 form a coherent track. Most mass-spectrometry talks cluster here.",
    "**Long-term monitoring:** Follow the method tag in the abstracts rather than session titles; it threads across lakes, climate, biodiversity, and the cryosphere.",
    "**Equity and Indigenous knowledge:** EP013 (Two-Eyed Seeing), WS05, WS08, SS079, and AV001 form a coherent, week-long arc.",
    "**Scale-bridging:** The nine \"multiple stressors\" and \"bridging the gap\" sessions are the map, anchored by SS082 (mesocosms) and SS070 (modelling/forecasts).",
    "**Cold and dry systems:** SS011, SS013, SS015, and SS019 function as a single conversation about climate-driven change, split across four sessions.",
]:
    p(doc, item, style="List Bullet")

doc.add_heading("Notes for Future Programmes", level=2)
p(doc, "A few signals from the data that future programmes might find useful, offered with no insider knowledge of why the current schedule looks the way it does:")
for item in [
    "**ECR infrastructure.** Six dedicated early-career sessions in five days is real structural support, and worth carrying forward.",
    "**Adjacent scheduling.** Tribute sessions scheduled adjacent to early-career tracks would let attendees walk straight from celebrating the past into the conversations that extend it.",
    "**Pacing.** The Friday-density signal could be eased over time by peak-day spacing, longer breaks between parallel blocks, or hybrid options for clash relief.",
    "**Bridge-session visibility.** Cross-disciplinary sessions like EP013 and AV001 carry a lot in the tag network for their size. Cross-track signposting would help them be easier to find.",
]:
    p(doc, item, style="List Bullet")

doc.add_heading("Final Thoughts", level=2)
p(doc, "The data points to a community crossing disciplinary boundaries, scaling its methods, pairing legacy with youth, and building equity into its scientific core. Have a brilliant time in Montreal.")

embed_chart(
    doc,
    REPO / "output" / "charts" / "meeting_highlights_figure1.png",
    "Figure 1. ",
    "Disciplinary distribution of all 1,455 scheduled presentations at ASLO-SIL 2026, by session-level frame. (Keyword classifier on session names, session descriptions, and presentation titles; multi-label. Columns overlap; percentages do not sum to 100%.) Approximately 93% of presentations carry at least one of the nine frame tags. Source: scraped from the public ASLO-SIL 2026 schedule, audited 11 May 2026.",
)
embed_chart(
    doc,
    REPO / "output" / "charts" / "freshwater_share_vs_participation_share.png",
    "Figure 2. ",
    "Country freshwater share (World Bank ER.H2O.INTR.K3) versus country share of presentations for which a country could be detected. Top-20 freshwater-stock countries shown. Japan is closest to parity; Brazil holds about 13% of the world's renewable freshwater and contributes about 1.3% of the programme.",
)

doc.add_paragraph()
dm = doc.add_paragraph()
r1 = dm.add_run("Data and methods: "); r1.italic = True; r1.bold = True
r2 = dm.add_run("Scraped from the public ASLO-SIL 2026 schedule and audited 11 May 2026. Disciplinary, method, and problem tags assigned by a multi-label keyword classifier. Country detection reaches approx. 78% coverage. Gender inference utilises the European-biased gender-guesser database (12% unclassified). Full inventory, classifier code, network/matrix visualisations, and reproducible analysis available at github.com/erika-freeman/aslo2026-conference-landscape.")
r2.italic = True

byline = doc.add_paragraph()
br = byline.add_run("Erika C. Freeman · IGB Berlin · SS050B \"Ecological Significance of Dissolved Organic Matter\" · 15 May 2026 · Palais des congrès de Montréal.")
br.italic = True

DOCX_NAME = "Freeman_ASLO-SIL_2026_Programme_Analysis_v2.docx"
PATH = PKG / DOCX_NAME
doc.save(str(PATH))
doc.save(str(BUNDLE / "04_Manuscript_Freeman_1100w.docx"))

# ===========================================================================
# Stage figures and watershed
# ===========================================================================
shutil.copy(WS_ARCHIVE_PNG, BUNDLE / "00_OpeningVisual_watershed.png")
shutil.copy(WS_ARCHIVE_TIF, BUNDLE / "00_OpeningVisual_watershed_CMYK_300dpi.tif")
shutil.copy(REPO / "output" / "charts" / "meeting_highlights_figure1.png",
            BUNDLE / "02_Figure1_disciplinary_distribution.png")
shutil.copy(PKG / "Freeman_MeetingHighlights_Figure1.tif",
            BUNDLE / "02_Figure1_disciplinary_distribution_CMYK_300dpi.tif")
shutil.copy(REPO / "output" / "charts" / "freshwater_share_vs_participation_share.png",
            BUNDLE / "03_Figure2_freshwater_vs_participation.png")

if ws_embed.exists():
    ws_embed.unlink()

# ===========================================================================
# Inquiry email v16 — same softened tone
# ===========================================================================
inquiry_v16 = """TO:       lobulletin-editor@aslo.org
FROM:     erika.freeman@igb-berlin.de
SUBJECT:  Presubmission inquiry: Meeting Highlights on ASLO-SIL 2026

ATTACHMENTS:
  1. 00_OpeningVisual_watershed.png
     (illustration, opening visual)
  2. 02_Figure1_disciplinary_distribution.png
     (disciplinary distribution of all 1,455 talks)
  3. 03_Figure2_freshwater_vs_participation.png
     (country freshwater share vs. programme share)
  4. 04_Manuscript_Freeman_1100w.docx
     (full draft, ~1,100 words, opening visual + two figures)

---

Dear Dr. Falkenberg,

I would like to pitch a Meeting Highlights piece for the L&O Bulletin, based on ASLO-SIL 2026 in Montreal this week.

Title: "ASLO-SIL 2026 in Numbers: Trends, Gaps, and Takeaways from the Programme."

I read the full programme before flying. 1,455 scheduled presentations across 308 session items, audited against the public schedule on 11 May 2026. The piece reads the schedule for what the data suggests about where the field is and some of the directions it appears to be moving in: which methods are doing the work, which conversations are routing through the schedule, where the programme looks stretched, and where some of the clearer opportunities for invitation sit.

A few of the load-bearing observations, each anchored to the underlying data:

  • The most-tagged single method in the 2026 programme is long-term monitoring (about 9% of talks), threading through about ten distinct problem clusters. AI and machine learning are present (about 2.4%) but spread thin: visible across many problems at low intensity in each.

  • Indigenous knowledge, equity, community-led, and citizen-science framings appear in 58 talks across 14 session items. In the programme's tag network they sit alongside rivers, estuaries, fisheries, food webs, and DOM as bridge topics connecting otherwise distinct sub-conversations. EP013 "Two-Eyed Seeing" is programmed as core science, not an adjunct.

  • Three career-tribute sessions (Mike Pace EP001, Jim Cotner EP004, Jim Elser EP012) appear in the same five-day programme as six dedicated early-career sessions. Whether intentional or not, the programme structurally pairs legacy and cohort.

  • The programme quietly shows a strong gender-balance signal at both presenter and organiser levels: approximately 53% female-inferred among classifiable presenter names and 56% among classifiable organiser names, using a European-biased name classifier that left about 12% of names unclassified.

  • The geography of participation points to one of the clearer opportunities the programme surfaces. US plus Canada account for 54% of presentations for which a country could be detected. Brazil (13% of global freshwater, 1.3% of the programme), Russia (10%), Colombia (5%), Indonesia (5%), Peru (4%), India (3%), and Myanmar (2%) together hold roughly 42% of the world's renewable freshwater and are under-represented or not detected. The piece treats this as a map for future collaboration, naming both programmatic levers (travel pathways, hybrid options, regional hubs) and deeper levers (international authorship valuation, funding-agency structures, language inclusivity in publishing).

The piece also includes a "Navigating the Meeting" reader guide and a short "Notes for Future Programmes" section, the latter offered as observations from a member, not advice from outside the schedule.

The draft is ~1,100 words, attached as a Word file with line numbers and page numbers for review. Every quantitative claim is traceable to a corresponding entry in the repository's data tables.

I am presenting at SS050B "Ecological Significance of DOM" on 15 May, so the piece is written as a participant-observer.

Would this fit the Bulletin? I am happy to revise to spec.

With thanks,

Erika

Dr. Erika C. Freeman
Group Leader, ABC Lab
Leibniz Institute of Freshwater Ecology and Inland Fisheries (IGB)
Müggelseedamm 310, 12587 Berlin, Germany
erika.freeman@igb-berlin.de  ·  ORCID 0000-0001-7161-6038
"""
(PKG / "presubmission_inquiry_email_v16.txt").write_text(inquiry_v16, encoding="utf-8")
(BUNDLE / "01_inquiry_email.txt").write_text(inquiry_v16, encoding="utf-8")

# ===========================================================================
# README
# ===========================================================================
readme = """SUBMISSION BUNDLE v11  |  L&O Bulletin presubmission inquiry
=============================================================
Recipient   : Dr. Laura Falkenberg, L&O Bulletin Editor
Address     : lobulletin-editor@aslo.org
From        : Erika C. Freeman (IGB Berlin)
Bundle date : 12 May 2026
Manuscript  : v18 (~1,100 words)
Inquiry     : v16

WHAT CHANGED SINCE BUNDLE v10
-----------------------------
Tone pass. Same content, softer voice.
  - Superlatives reduced or removed throughout the body.
  - Imperative "attendees should…" reframed as observation.
  - The recommendations section is now "Notes for Future Programmes,"
    offered with an explicit acknowledgment that the author has no
    insider view of why the schedule is the way it is. Each bullet
    recast from imperative to observation.
  - "Final Thoughts" depuffed: "The data points to a community
    crossing…" rather than "The data reveals a community that is
    actively crossing…"

Manuscript formatting defaults (Times New Roman 12pt, 1.5 spacing,
continuous line numbers, page numbers in footer) preserved.

WHAT TO DO
----------
1. Open 01_inquiry_email.txt. Copy the body (below the "---" divider) into
   a new email to lobulletin-editor@aslo.org. Use the SUBJECT line at top.

2. Attach all four files:
       00_OpeningVisual_watershed.png
       02_Figure1_disciplinary_distribution.png
       03_Figure2_freshwater_vs_participation.png
       04_Manuscript_Freeman_1100w.docx

3. Send.

WHAT'S IN THE BUNDLE
--------------------
00_README.txt                                   This file.
01_inquiry_email.txt                            Email body (v16).
00_OpeningVisual_watershed.png                  Opening visual.
00_OpeningVisual_watershed_CMYK_300dpi.tif      Print-ready.
02_Figure1_disciplinary_distribution.png        Data figure 1.
02_Figure1_disciplinary_distribution_CMYK_300dpi.tif
03_Figure2_freshwater_vs_participation.png      Data figure 2.
04_Manuscript_Freeman_1100w.docx                Full draft.

Canonical archive (submission_package/):
    Freeman_ASLO-SIL_2026_Programme_Analysis_v2.docx
    presubmission_inquiry_email_v16.txt

Build script:
.../_GitHubRepo/aslo2026-conference-landscape/scripts/24_build_submission_bundle_v11.py
"""
(BUNDLE / "00_README.txt").write_text(readme, encoding="utf-8")

# ===========================================================================
# Sanity check
# ===========================================================================
forbidden = ["commission", "hand-drawn", "ultimate", "vastly", "massive freshwater",
             "remarkably close", "completely undetected", "impressive breadth",
             "vertebrae", "backbone"]
for label, fp in [
    ("Manuscript", BUNDLE / "04_Manuscript_Freeman_1100w.docx"),
    ("Inquiry email", BUNDLE / "01_inquiry_email.txt"),
    ("README", BUNDLE / "00_README.txt"),
]:
    if fp.suffix == ".docx":
        text = "\n".join(par.text for par in Document(str(fp)).paragraphs)
    else:
        text = fp.read_text(encoding="utf-8")
    hits = [w for w in forbidden if w.lower() in text.lower()]
    print(("OK: " if not hits else "WARN: ") + label + (" clean" if not hits else " contains " + str(hits)))

print()
print("Bundle assembled at: {}".format(BUNDLE))
print()
print("Contents:")
for f in sorted(BUNDLE.iterdir()):
    print("  {:>10} bytes  {}".format(f.stat().st_size, f.name))

def body_wc(p):
    return sum(len(par.text.split()) for par in Document(str(p)).paragraphs)
print()
print("Manuscript docx total word count: {}".format(body_wc(BUNDLE / "04_Manuscript_Freeman_1100w.docx")))
